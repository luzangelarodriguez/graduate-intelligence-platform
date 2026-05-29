from __future__ import annotations

from pathlib import Path

import pytest

from microcurriculum_engine.evaluation.scoring import score_microcurriculum
from microcurriculum_engine.ingestion.document_loader import load_document
from microcurriculum_engine.matching.market_matching import compare_microcurriculum_to_market
from microcurriculum_engine.normalization.skill_extractor import extract_microcurriculum_skills
from microcurriculum_engine.parsing.academic_parser import parse_microcurriculum
from microcurriculum_engine.pipelines.process_microcurriculum import process_microcurriculum
from microcurriculum_engine.recommendations.recommendation_engine import generate_recommendations


SAMPLE_TEXT = """
Programa: Especialización en Ingeniería de Software
Asignatura: Desarrollo de Aplicaciones Web
Semestre: 1
Créditos: 3

Competencias:
- Diseñar servicios backend con Python y SQL.
- Implementar interfaces con React y JavaScript.

Resultados de aprendizaje:
- Construye APIs RESTful y despliega aplicaciones con Docker.

Contenidos:
Python, SQL, React, Docker, REST API, arquitectura de software.

Metodologías:
Scrum, Agile, trabajo en equipo y pensamiento crítico.

Herramientas:
AWS, Docker, GitHub Actions.
"""


def test_text_extraction_from_txt(tmp_path: Path) -> None:
    source = tmp_path / "microcurriculo.txt"
    source.write_text(SAMPLE_TEXT, encoding="utf-8")
    document = load_document(source, persist_original=False)
    assert "Python" in document.clean_text
    assert document.extraction_method == "plain_text"


def test_text_extraction_from_docx_tables(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    import docx

    source = tmp_path / "microcurriculo.docx"
    document = docx.Document()
    document.add_paragraph("Programa academico")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Aprendizaje automatico"
    table.cell(0, 1).text = "Machine learning"
    document.save(source)

    loaded = load_document(source, persist_original=False)

    assert loaded.extraction_method == "python-docx"
    assert "Aprendizaje automatico" in loaded.clean_text
    assert "Machine learning" in loaded.clean_text


def test_academic_parser_detects_core_fields() -> None:
    parsed = parse_microcurriculum(SAMPLE_TEXT)
    assert parsed.programa == "Especialización en Ingeniería de Software"
    assert parsed.asignatura == "Desarrollo de Aplicaciones Web"
    assert parsed.creditos == "3"
    assert parsed.competencias
    assert parsed.resultados_aprendizaje


def test_skill_detection_and_normalization() -> None:
    extracted = extract_microcurriculum_skills(SAMPLE_TEXT, title="Desarrollo de Aplicaciones Web")
    skill_names = {skill["skill_normalized"] for skill in extracted["skills"]}
    assert {"python", "sql", "docker", "react"}.issubset(skill_names)
    assert extracted["domain_prediction"]["domain"] in {"ti", "analitica"}


def test_gap_recommendation_and_scoring() -> None:
    comparison = compare_microcurriculum_to_market(
        ["python", "sql"],
        domain="ti",
        market_skills=["python", "sql", "docker", "react"],
    )
    recommendations = generate_recommendations(domain="ti", comparison=comparison)
    scores = score_microcurriculum(
        comparison=comparison,
        skills_count=2,
        competencies_count=3,
        recommendations_count=len(recommendations),
    )
    assert comparison.missing_skills == ["docker", "react"]
    assert any("Docker" in item["title"] or "docker" in item["title"] for item in recommendations)
    assert 0 <= scores["pertinencia_curricular"] <= 1


def test_pipeline_processes_microcurriculum_without_database(tmp_path: Path) -> None:
    source = tmp_path / "microcurriculo.txt"
    source.write_text(SAMPLE_TEXT, encoding="utf-8")
    result = process_microcurriculum(
        source,
        persist=False,
        market_skills=["python", "sql", "docker", "react", "devops"],
    )
    assert result["document"]["clean_text"]
    assert result["skills"]
    assert "devops" in result["gaps"]["missing_skills"]
    assert result["recommendations"]
    assert result["scores"]["alineacion_laboral"] >= 0
    assert result["embeddings"]["documento"]["dimensions"] > 0
