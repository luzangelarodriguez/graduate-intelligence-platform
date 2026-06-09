#!/usr/bin/env python3
"""
scripts/load_microcurriculos.py

Parse all .docx microcurriculum files from storage/test_microcurriculos/,
extract structured academic data, and load into the DB.

Usage:
    python scripts/load_microcurriculos.py --preview   # show what would be inserted
    python scripts/load_microcurriculos.py --execute   # confirm & insert

DB connection: RAILWAY_DATABASE_URL from .env.local (root of repo).
NO data is written without --execute.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sys
import unicodedata
from pathlib import Path
from typing import Any

# ── paths ─────────────────────────────────────────────────────────────────────

REPO_ROOT = Path(__file__).resolve().parent.parent
STORAGE_DIR = REPO_ROOT / "storage" / "test_microcurriculos"
OUTPUTS_DIR = REPO_ROOT / "outputs"
ENV_LOCAL = REPO_ROOT / ".env.local"

# ── domain mapping ─────────────────────────────────────────────────────────────

# programa name fragment  →  domain_key
DOMAIN_MAP: dict[str, str] = {
    "visual analytics": "data_analytics",
    "big data": "data_analytics",
    "inteligencia artificial": "artificial_intelligence",
    "machine learning": "artificial_intelligence",
    "criminolog": "criminology",
    "victimolog": "criminology",
    "gerencia": "business",
    "administracion": "business",
    "gestion de proyectos": "business",
    "direccion": "business",
    "derecho": "law",
    "educacion": "education",
    "ambiental": "health",
    "salud": "health",
}


def infer_domain(programa: str) -> str:
    norm = _normalize_text(programa)
    for fragment, domain in DOMAIN_MAP.items():
        if fragment in norm:
            return domain
    return "general"


# ── skill extraction regexes ───────────────────────────────────────────────────

_RE_PLATFORMS = re.compile(
    r"\b(AWS|Azure|GCP|Google\s*Cloud|Databricks|Snowflake|BigQuery|"
    r"Power\s*Apps|Power\s*Automate|Power\s*Platform|MLflow|Airflow|"
    r"Cloudera|Hortonworks|Vertica|Redshift)\b",
    re.IGNORECASE,
)

_RE_TOOLS = re.compile(
    r"\b(Power\s*BI|Tableau|Looker|QlikSense|QlikView|Excel|KNIME|"
    r"RStudio|Jupyter|VS\s*Code|PyCharm|SPSS|SAS|Minitab|"
    r"Docker|Kubernetes|Git|GitHub|GitLab|Jenkins|dbt|"
    r"Pentaho|Talend|SSIS|Informatica|MicroStrategy|SAP)\b",
    re.IGNORECASE,
)

_RE_TECH = re.compile(
    r"\b(Python|R\b|SQL|NoSQL|Spark|Hadoop|Kafka|Flink|"
    r"TensorFlow|PyTorch|Keras|Scikit.?learn|XGBoost|LightGBM|"
    r"MongoDB|PostgreSQL|MySQL|SQLite|Cassandra|Redis|Elasticsearch|"
    r"OLAP|OLTP|ETL|ELT|Data\s*Warehouse|Data\s*Lake|Data\s*Mart|"
    r"Machine\s*Learning|Deep\s*Learning|NLP|Computer\s*Vision|"
    r"Inteligencia\s*Artificial|Aprendizaje\s*Autom[aá]tico|"
    r"Regresi[oó]n|Clasificaci[oó]n|Clustering|Random\s*Forest|"
    r"Dashboard|Reporting|API\s*REST|Microservicios|"
    r"Visualizaci[oó]n\s*de\s*datos|An[aá]lisis\s*de\s*datos)\b",
    re.IGNORECASE,
)

_RE_TRANSVERSAL = re.compile(
    r"\b(Liderazgo|Comunicaci[oó]n\s+efectiva|Comunicaci[oó]n|"
    r"Trabajo\s+en\s+equipo|Pensamiento\s+cr[ií]tico|"
    r"Resoluci[oó]n\s+de\s+problemas|Toma\s+de\s+decisiones|"
    r"Gesti[oó]n\s+del\s+cambio|Gesti[oó]n|Planeaci[oó]n|"
    r"Emprendimiento|Innovaci[oó]n|[EÉ]tica\s+profesional|[EÉ]tica|"
    r"Responsabilidad\s+social|Responsabilidad|Adaptabilidad|"
    r"Creatividad|Negociaci[oó]n|Presentaci[oó]n)\b",
    re.IGNORECASE,
)

_RE_METHODS = re.compile(
    r"\b(Scrum|Agile|CRISP.?DM|Design\s*Thinking|Lean|Six\s*Sigma|"
    r"PMBOK|PRINCE2|Kanban|DevOps|DataOps|MLOps|"
    r"Regresi[oó]n\s+lineal|Regresi[oó]n\s+log[ií]stica|"
    r"An[aá]lisis\s+factorial|Series\s+de\s+tiempo|"
    r"Validaci[oó]n\s+cruzada|Hiperpar[aá]metros|"
    r"An[aá]lisis\s+de\s+componentes\s+principales|PCA|"
    r"Balanceo\s+de\s+clases|Overfitting|Underfitting)\b",
    re.IGNORECASE,
)


def _normalize_text(text: str) -> str:
    nfkd = unicodedata.normalize("NFKD", text.lower())
    return "".join(c for c in nfkd if not unicodedata.combining(c)).strip()


def _title(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip()).title()


def extract_skills(text: str) -> dict[str, list[str]]:
    buckets: dict[str, list[str]] = {
        "tecnologia": [],
        "skill_tecnica": [],
        "herramienta": [],
        "plataforma": [],
        "skill_transversal": [],
        "metodologia": [],
    }
    seen: set[str] = set()

    def _add(bucket: str, raw: str) -> None:
        key = _normalize_text(raw)
        if key not in seen:
            seen.add(key)
            buckets[bucket].append(_title(raw))

    for m in _RE_PLATFORMS.finditer(text):
        _add("plataforma", m.group(0))
    for m in _RE_TOOLS.finditer(text):
        _add("herramienta", m.group(0))
    for m in _RE_TECH.finditer(text):
        raw = m.group(0)
        key = _normalize_text(raw)
        if key not in seen:
            seen.add(key)
            buckets["tecnologia"].append(_title(raw))
    for m in _RE_TRANSVERSAL.finditer(text):
        _add("skill_transversal", m.group(0))
    for m in _RE_METHODS.finditer(text):
        _add("metodologia", m.group(0))

    return {k: v for k, v in buckets.items() if v}


# ── docx helpers ───────────────────────────────────────────────────────────────

def _cell(table, row: int, col: int = 0) -> str:
    try:
        return table.rows[row].cells[col].text.strip()
    except IndexError:
        return ""


def _table_for(tables, keyword: str):
    for t in tables:
        if _cell(t, 0) and keyword.upper() in _cell(t, 0).upper():
            return t
    return None


def _table_all_text(table) -> str:
    parts: list[str] = []
    for row in table.rows[1:]:
        for cell in row.cells:
            txt = cell.text.strip()
            if txt:
                parts.append(txt)
    return "\n".join(parts)


# ── parse a standard 10-table microcurriculum docx ────────────────────────────

def parse_standard_docx(path: Path) -> dict[str, Any] | None:
    from docx import Document  # type: ignore

    doc = Document(str(path))
    tables = doc.tables
    if len(tables) < 6:
        return None

    programa = _cell(_table_for(tables, "PROGRAMA ACADÉMICO") or tables[0], 1)
    asignatura = _cell(_table_for(tables, "DENOMINACIÓN") or tables[1], 1)
    if not programa or not asignatura:
        return None

    desc_table = _table_for(tables, "DESCRIPCIÓN")
    ra_table = _table_for(tables, "RESULTADOS DE APRENDIZAJE")
    media_table = _table_for(tables, "MEDIOS EDUCATIVOS")

    descripcion = _table_all_text(desc_table) if desc_table else ""

    # Split RA table: first body cell = learning outcomes, second = thematic content
    resultados_text = ""
    contenido_tematico = ""
    if ra_table:
        body_cells = [
            c.text.strip()
            for row in ra_table.rows[1:]
            for c in row.cells
            if c.text.strip()
        ]
        if len(body_cells) >= 1:
            resultados_text = body_cells[0]
        if len(body_cells) >= 2:
            contenido_tematico = body_cells[1]

    medios_text = _table_all_text(media_table) if media_table else ""

    # Learning outcomes: lines that start with a verb or "RA"
    resultados: list[str] = []
    for line in resultados_text.splitlines():
        line = line.strip().rstrip('"')
        if len(line) > 20 and re.match(r"^[A-ZÁÉÍÓÚÑ]", line):
            if not re.match(r"^(Tema|Contenido|Nota|Bibliograf|\*)", line, re.I):
                resultados.append(line)

    # Thematic content: extract topics
    temas: list[str] = []
    for line in contenido_tematico.splitlines():
        line = line.strip()
        if re.match(r"^Tema\s+\d", line, re.I):
            temas.append(line)

    full_text = "\n".join([descripcion, resultados_text, contenido_tematico, medios_text])
    skills = extract_skills(full_text)

    doc_hash = hashlib.md5(path.read_bytes()).hexdigest()

    return {
        "programa": programa,
        "asignatura": asignatura,
        "descripcion": descripcion[:2000],
        "resultados_aprendizaje": resultados[:15],
        "contenido_tematico": temas or [contenido_tematico[:500]],
        "herramientas_recursos": medios_text[:1000],
        "skills": skills,
        "source_document": path.name,
        "source_path": str(path.relative_to(REPO_ROOT)),
        "document_hash": doc_hash,
        "domain_key": infer_domain(programa),
    }


# ── parse criminología program docx (paragraph format) ───────────────────────

def parse_criminologia_docx(path: Path) -> list[dict[str, Any]]:
    from docx import Document  # type: ignore

    doc = Document(str(path))
    programa = "Especialización en Criminología y Victimología"

    # Find RA table (first table with "Resultado" in header)
    ra_table = None
    for t in doc.tables:
        if "RESULTADO" in _cell(t, 0).upper():
            ra_table = t
            break
    if ra_table is None:
        return []

    # Parse asignaturas and their RAs from merged-cell table
    # Pattern: alternating RA rows and asignatura name rows
    asig_ras: dict[str, list[str]] = {}
    current = ""
    for row in ra_table.rows[1:]:
        unique = list(dict.fromkeys(c.text.strip() for c in row.cells if c.text.strip()))
        if not unique:
            continue
        # Asignatura name: appears in rightmost unique cell, doesn't start with "RA"
        last = unique[-1]
        if (last and not last.startswith("RA") and not last.startswith("•")
                and len(last) > 5 and len(last) < 120 and last != current):
            # Only accept if it's clearly an asignatura (not a RA statement)
            if not re.match(r"^(Analizar|Examinar|Explicar|Desarrollar|Aplicar|Identificar)", last):
                current = last
                if current not in asig_ras:
                    asig_ras[current] = []
        first = unique[0]
        if first.startswith("RA") and current:
            asig_ras[current].append(first[:300])

    # Program-level skills (from all paragraph text)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    skills = extract_skills(full_text)
    doc_hash = hashlib.md5(path.read_bytes()).hexdigest()

    results: list[dict[str, Any]] = []
    for asig, ras in asig_ras.items():
        if not asig or asig.lower() in ("resultados de aprendizaje", "asignatura"):
            continue
        results.append({
            "programa": programa,
            "asignatura": asig,
            "descripcion": "",
            "resultados_aprendizaje": ras[:10],
            "contenido_tematico": [],
            "herramientas_recursos": "",
            "skills": skills,
            "source_document": path.name,
            "source_path": str(path.relative_to(REPO_ROOT)),
            "document_hash": doc_hash + f":{asig[:20]}",
            "domain_key": "criminology",
        })
    return results


# ── discover all docx ─────────────────────────────────────────────────────────

def discover_all() -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for docx_path in sorted(STORAGE_DIR.rglob("*.docx")):
        folder = docx_path.relative_to(STORAGE_DIR).parts[0]
        if "Criminolog" in folder or "criminolog" in folder:
            records.extend(parse_criminologia_docx(docx_path))
        else:
            parsed = parse_standard_docx(docx_path)
            if parsed:
                records.append(parsed)
    return records


# ── DB helpers ────────────────────────────────────────────────────────────────

def _load_env() -> None:
    if ENV_LOCAL.exists():
        try:
            from dotenv import load_dotenv  # type: ignore
            load_dotenv(ENV_LOCAL, override=True)
        except ImportError:
            # manual parse
            for line in ENV_LOCAL.read_text().splitlines():
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    k, _, v = line.partition("=")
                    os.environ.setdefault(k.strip(), v.strip().strip('"').strip("'"))


def get_conn():
    _load_env()
    url = os.environ.get("RAILWAY_DATABASE_URL") or os.environ.get("DATABASE_URL")
    if not url:
        raise RuntimeError(
            "RAILWAY_DATABASE_URL not set.\n"
            "Create .env.local in the repo root with:\n"
            "  RAILWAY_DATABASE_URL=postgresql://user:pass@host:port/db"
        )
    import psycopg2  # type: ignore
    return psycopg2.connect(url)


def fetch_especializaciones(conn) -> list[dict]:
    with conn.cursor() as cur:
        cur.execute("SELECT id, nombre FROM especializaciones ORDER BY id")
        return [{"id": r[0], "nombre": r[1]} for r in cur.fetchall()]


# Explicit overrides: normalized program name fragment → especializacion_id
_EXPLICIT_ESP_MAP: dict[str, int] = {
    # Criminología / Psicología Criminal / Victimología → id=108
    "psicologia criminal": 108,
    "criminologia": 108,
    "victimologia": 108,
    "criminolog": 108,
    "victimolog": 108,
    # Administración de Empresas → id=82 (Alta Gerencia)
    "administracion de empresas": 82,
    "alta gerencia": 82,
}


def match_esp_id(programa: str, esps: list[dict]) -> int | None:
    prog_n = _normalize_text(programa)

    # 1. Explicit map takes priority
    for fragment, eid in _EXPLICIT_ESP_MAP.items():
        if fragment in prog_n:
            return eid

    # 2. Fuzzy match — collect all candidates, prefer highest id
    candidates: list[int] = []
    for e in esps:
        e_n = _normalize_text(e["nombre"])
        if e_n in prog_n or prog_n in e_n:
            candidates.append(e["id"])

    return max(candidates) if candidates else None


def fetch_existing_microcurriculos(conn) -> set[str]:
    """Return set of 'programa|asignatura' already in DB (normalized)."""
    with conn.cursor() as cur:
        cur.execute("SELECT programa, asignatura FROM microcurriculos WHERE programa IS NOT NULL")
        return {
            _normalize_text(f"{r[0]}|{r[1]}")
            for r in cur.fetchall()
            if r[0] and r[1]
        }


def fetch_existing_hashes(conn) -> set[str]:
    with conn.cursor() as cur:
        cur.execute("SELECT document_hash FROM microcurriculos WHERE document_hash IS NOT NULL")
        return {r[0] for r in cur.fetchall()}


def _has_column(conn, table: str, column: str) -> bool:
    with conn.cursor() as cur:
        cur.execute(
            """
            SELECT 1 FROM information_schema.columns
            WHERE table_schema='public' AND table_name=%s AND column_name=%s
            """,
            (table, column),
        )
        return cur.fetchone() is not None


def insert_record(conn, rec: dict, esp_id: int | None) -> int:
    with conn.cursor() as cur:
        cur.execute(
            """
            INSERT INTO microcurriculos
                (programa, asignatura, semestre, source_document,
                 document_hash, clean_text, detected_domain, lineage,
                 specialization_id, specialization_name)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            ON CONFLICT (document_hash) DO NOTHING
            RETURNING id
            """,
            (
                rec["programa"],
                rec["asignatura"],
                "",
                rec["source_document"],
                rec["document_hash"],
                "\n\n".join([
                    rec["descripcion"],
                    "\n".join(rec["resultados_aprendizaje"]),
                    "\n".join(rec["contenido_tematico"]),
                ]).strip(),
                rec["domain_key"],
                json.dumps({
                    "source_path": rec["source_path"],
                    "domain_key": rec["domain_key"],
                }),
                esp_id,
                rec["programa"] if esp_id else None,
            ),
        )
        row = cur.fetchone()
        if row is None:
            # conflict — fetch existing id
            cur.execute(
                "SELECT id FROM microcurriculos WHERE document_hash = %s",
                (rec["document_hash"],),
            )
            return cur.fetchone()[0]
        return row[0]


def insert_skills(conn, micro_id: int, rec: dict) -> int:
    rows_inserted = 0
    with conn.cursor() as cur:
        for tipo, skills in rec["skills"].items():
            for skill in skills:
                try:
                    cur.execute(
                        """
                        INSERT INTO microcurriculo_skills
                            (microcurriculo_id, skill_original, skill_normalized,
                             skill_domain, tipo_skill, confidence_score, source_document)
                        VALUES (%s, %s, %s, %s, %s, %s, %s)
                        ON CONFLICT (microcurriculo_id, skill_normalized, tipo_skill) DO NOTHING
                        """,
                        (
                            micro_id,
                            skill,
                            skill,
                            rec["domain_key"],
                            tipo,
                            0.70,
                            rec["source_document"],
                        ),
                    )
                    rows_inserted += cur.rowcount
                except Exception:
                    pass  # ignore individual skill conflicts
    return rows_inserted


def upsert_domain_mapping(conn, esp_id: int, programa: str, domain_key: str) -> None:
    """Update detected_domain on especializaciones if column exists."""
    if not _has_column(conn, "especializaciones", "detected_domain"):
        return
    with conn.cursor() as cur:
        cur.execute(
            """
            UPDATE especializaciones
               SET detected_domain = %s, updated_at = now()
             WHERE id = %s AND (detected_domain IS NULL OR detected_domain = '')
            """,
            (domain_key, esp_id),
        )


# ── preview renderer ──────────────────────────────────────────────────────────

def print_preview(records: list[dict], existing_keys: set[str], esps: list[dict]) -> None:
    by_prog: dict[str, list[dict]] = {}
    for r in records:
        by_prog.setdefault(r["programa"], []).append(r)

    total_skills = sum(sum(len(v) for v in r["skills"].values()) for r in records)
    new_records = [
        r for r in records
        if _normalize_text(f"{r['programa']}|{r['asignatura']}") not in existing_keys
    ]
    dup_count = len(records) - len(new_records)

    print()
    print("╔══════════════════════════════════════════════════════════╗")
    print("║           PREVIEW — microcurriculos a insertar           ║")
    print("╚══════════════════════════════════════════════════════════╝")
    print(f"  Programas encontrados   : {len(by_prog)}")
    print(f"  Asignaturas totales     : {len(records)}")
    print(f"  Skills extraídos        : {total_skills}")
    print(f"  Ya en DB (duplicados)   : {dup_count}")
    print(f"  A insertar              : {len(new_records)}")
    print()

    for prog, recs in sorted(by_prog.items()):
        esp_id = match_esp_id(prog, esps)
        domain = recs[0]["domain_key"]
        print(f"  📚 {prog}")
        print(f"     especializacion_id = {esp_id}   domain_key = {domain}")
        for r in recs:
            key = _normalize_text(f"{r['programa']}|{r['asignatura']}")
            status = "✓ ya existe" if key in existing_keys else "+ NUEVO"
            n_skills = sum(len(v) for v in r["skills"].values())
            n_ra = len(r["resultados_aprendizaje"])
            print(f"     [{status}]  {r['asignatura']}")
            print(f"              RA={n_ra}  skills={n_skills}  src={r['source_document'][:50]}")
        print()


# ── report writer ─────────────────────────────────────────────────────────────

def write_report(
    records: list[dict],
    existing_keys: set[str],
    esps: list[dict],
    db_available: bool,
    inserted: list[dict] | None,
) -> None:
    OUTPUTS_DIR.mkdir(exist_ok=True)
    report_path = OUTPUTS_DIR / "microcurriculos_load_report.md"

    new_records = [
        r for r in records
        if _normalize_text(f"{r['programa']}|{r['asignatura']}") not in existing_keys
    ]

    lines: list[str] = []
    lines.append("# Reporte: Carga de Microcurrículos")
    lines.append(f"**Fecha:** 2026-06-08  ")
    lines.append(f"**Script:** `scripts/load_microcurriculos.py`  ")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Resumen")
    lines.append("")
    lines.append(f"| Métrica | Valor |")
    lines.append(f"|---------|-------|")
    lines.append(f"| Archivos .docx encontrados | {len(set(r['source_document'] for r in records))} |")
    lines.append(f"| Asignaturas extraídas | {len(records)} |")
    lines.append(f"| Skills extraídos | {sum(sum(len(v) for v in r['skills'].values()) for r in records)} |")
    lines.append(f"| DB disponible | {'✅ Sí' if db_available else '❌ No (sin .env.local)'} |")
    lines.append(f"| Ya en DB (duplicados) | {len(records) - len(new_records)} |")
    lines.append(f"| A insertar (nuevos) | {len(new_records)} |")
    if inserted is not None:
        lines.append(f"| Insertados en esta ejecución | {len(inserted)} |")
    lines.append("")

    lines.append("## Archivos encontrados")
    lines.append("")
    lines.append("| # | Archivo | Programa | Asignaturas |")
    lines.append("|---|---------|----------|------------|")
    seen_docs: dict[str, int] = {}
    for r in records:
        seen_docs[r["source_document"]] = seen_docs.get(r["source_document"], 0) + 1
    for i, (doc, count) in enumerate(sorted(seen_docs.items()), 1):
        prog = next((r["programa"] for r in records if r["source_document"] == doc), "")
        lines.append(f"| {i} | `{doc[:60]}` | {prog[:50]} | {count} |")
    lines.append("")

    lines.append("## Detalle por programa")
    lines.append("")
    by_prog: dict[str, list[dict]] = {}
    for r in records:
        by_prog.setdefault(r["programa"], []).append(r)

    for prog, recs in sorted(by_prog.items()):
        esp_id = match_esp_id(prog, esps) if esps else None
        domain = recs[0]["domain_key"]
        lines.append(f"### {prog}")
        lines.append(f"- `especializacion_id` = {esp_id}  ")
        lines.append(f"- `domain_key` = `{domain}`  ")
        lines.append(f"- Asignaturas: {len(recs)}")
        lines.append("")
        lines.append("| Asignatura | RA | Skills | Estado |")
        lines.append("|------------|-----|--------|--------|")
        for r in recs:
            key = _normalize_text(f"{r['programa']}|{r['asignatura']}")
            status = "Ya en DB" if key in existing_keys else "Nuevo"
            n_skills = sum(len(v) for v in r["skills"].values())
            lines.append(f"| {r['asignatura'][:60]} | {len(r['resultados_aprendizaje'])} | {n_skills} | {status} |")
        lines.append("")

    lines.append("## Mapeo de domain_key")
    lines.append("")
    lines.append("| Programa (fragmento) | domain_key |")
    lines.append("|----------------------|------------|")
    for frag, dk in DOMAIN_MAP.items():
        lines.append(f"| `{frag}` | `{dk}` |")
    lines.append("")

    lines.append("## Estructura de tablas usadas")
    lines.append("")
    lines.append("### microcurriculos")
    lines.append("```")
    lines.append("programa, asignatura, semestre, source_document, document_hash,")
    lines.append("clean_text, detected_domain, specialization_id, specialization_name, lineage")
    lines.append("```")
    lines.append("")
    lines.append("### microcurriculo_skills")
    lines.append("```")
    lines.append("microcurriculo_id, skill_original, skill_normalized,")
    lines.append("skill_domain, tipo_skill (tecnologia|skill_tecnica|herramienta|")
    lines.append("plataforma|skill_transversal|metodologia), confidence_score=0.70, source_document")
    lines.append("```")
    lines.append("")
    lines.append("### especializaciones (UPDATE detected_domain si columna existe)")
    lines.append("```")
    lines.append("UPDATE especializaciones SET detected_domain = <domain_key> WHERE id = <esp_id>")
    lines.append("```")
    lines.append("")

    lines.append("## Instrucciones de ejecución")
    lines.append("")
    lines.append("```bash")
    lines.append("# 1. Crear .env.local con la URL de producción")
    lines.append("echo 'RAILWAY_DATABASE_URL=postgresql://...' > .env.local")
    lines.append("")
    lines.append("# 2. Preview (sin tocar la DB)")
    lines.append("python scripts/load_microcurriculos.py --preview")
    lines.append("")
    lines.append("# 3. Ejecutar inserts (pide confirmación)")
    lines.append("python scripts/load_microcurriculos.py --execute")
    lines.append("```")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("*Reporte generado: 2026-06-08*")

    report_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"\n  Reporte escrito en: {report_path.relative_to(REPO_ROOT)}")


# ── main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--preview", action="store_true", help="Show preview without inserting")
    parser.add_argument("--execute", action="store_true", help="Insert after preview confirmation")
    args = parser.parse_args()

    if not args.preview and not args.execute:
        parser.print_help()
        print("\nHint: use --preview to see what would be inserted.")
        sys.exit(0)

    # 1. Parse docx
    print("[1/4] Parsing docx files from storage/test_microcurriculos/ …")
    records = discover_all()
    print(f"  {len(records)} microcurriculum records parsed from docx")

    # 2. Connect to DB (best effort)
    conn = None
    esps: list[dict] = []
    existing_keys: set[str] = set()
    db_available = False
    print("[2/4] Connecting to DB …")
    try:
        conn = get_conn()
        esps = fetch_especializaciones(conn)
        existing_keys = fetch_existing_microcurriculos(conn)
        db_available = True
        print(f"  DB reachable — {len(esps)} especializaciones, {len(existing_keys)} microcurriculos existing")
    except Exception as exc:
        print(f"  WARNING: DB not reachable — {exc}")
        print("  Continuing in offline mode (no duplicate detection).")

    # 3. Preview
    print("[3/4] Building preview …")
    print_preview(records, existing_keys, esps)

    new_records = [
        r for r in records
        if _normalize_text(f"{r['programa']}|{r['asignatura']}") not in existing_keys
    ]

    # 4. Execute
    inserted: list[dict] | None = None
    if args.execute:
        if not db_available:
            print("ERROR: Cannot execute — DB not reachable. Set RAILWAY_DATABASE_URL in .env.local")
            _write_report_only(records, existing_keys, esps, db_available)
            sys.exit(1)

        print(f"[4/4] Ready to insert {len(new_records)} records.")
        answer = input("  Confirm? [yes/no]: ").strip().lower()
        if answer not in ("yes", "y"):
            print("  Aborted.")
            sys.exit(0)

        inserted = []
        conn.autocommit = False
        try:
            for rec in new_records:
                esp_id = match_esp_id(rec["programa"], esps)
                micro_id = insert_record(conn, rec, esp_id)
                n_skills = insert_skills(conn, micro_id, rec)
                if esp_id:
                    upsert_domain_mapping(conn, esp_id, rec["programa"], rec["domain_key"])
                inserted.append({"asignatura": rec["asignatura"], "id": micro_id, "skills": n_skills})
                print(f"  ✓ id={micro_id}  {rec['asignatura']}  ({n_skills} skills)")
            conn.commit()
            print(f"\n  Committed. {len(inserted)} microcurriculos inserted.")
        except Exception as exc:
            conn.rollback()
            print(f"\n  ERROR — rolled back: {exc}")
            sys.exit(1)
        finally:
            conn.close()
    else:
        print("[4/4] Dry-run complete. Run with --execute to insert.")

    write_report(records, existing_keys, esps, db_available, inserted)


def _write_report_only(records, existing_keys, esps, db_available):
    write_report(records, existing_keys, esps, db_available, None)


if __name__ == "__main__":
    main()
