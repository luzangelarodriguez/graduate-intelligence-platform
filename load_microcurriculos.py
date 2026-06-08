#!/usr/bin/env python3
"""
load_microcurriculos.py — Parse docx microcurriculum files and load to DB.

Usage:
    python load_microcurriculos.py --dry-run          # preview only (default)
    python load_microcurriculos.py --execute          # insert into production DB
    python load_microcurriculos.py --dry-run --verbose

CRITICAL: --dry-run is the default. Use --execute only after reviewing the preview.
"""
from __future__ import annotations

import argparse
import json
import os
import re
import sys
import unicodedata
from pathlib import Path
from typing import Any

# ── constants ─────────────────────────────────────────────────────────────────

STORAGE_DIR = Path(__file__).parent / "storage" / "test_microcurriculos"

# Skill-extraction heuristics: keywords found in contenido temático / resultados
_TECH_KEYWORDS = re.compile(
    r"\b(Python|R\b|SQL|NoSQL|Spark|Hadoop|Kafka|TensorFlow|PyTorch|Keras|"
    r"Scikit.?learn|Power\s*BI|Tableau|Looker|QlikSense|Excel|KNIME|"
    r"AWS|Azure|GCP|Google\s*Cloud|Databricks|Snowflake|BigQuery|"
    r"MongoDB|PostgreSQL|MySQL|Cassandra|Redis|Elasticsearch|"
    r"Docker|Kubernetes|Git|GitHub|Airflow|dbt|MLflow|"
    r"Machine\s*Learning|Deep\s*Learning|NLP|Computer\s*Vision|"
    r"ETL|ELT|Data\s*Warehouse|Data\s*Lake|Data\s*Mart|"
    r"Regression|Clasificaci[oó]n|Clustering|Random\s*Forest|XGBoost|"
    r"Power\s*Apps|Power\s*Automate|Power\s*Platform|DAX|M\s+Query|"
    r"Inteligencia\s*Artificial|Aprendizaje\s*Autom[aá]tico|"
    r"Visualizaci[oó]n|Dashboard|Reporting|OLAP|OLTP)",
    re.IGNORECASE,
)

_TRANSVERSAL_KEYWORDS = re.compile(
    r"\b(Liderazgo|Comunicaci[oó]n|Trabajo\s*en\s*equipo|Pensamiento\s*cr[ií]tico|"
    r"Resoluci[oó]n\s*de\s*problemas|Toma\s*de\s*decisiones|Gesti[oó]n|"
    r"Planeaci[oó]n|Emprendimiento|Innovaci[oó]n|[EÉ]tica|Responsabilidad|"
    r"Adaptabilidad|Creatividad|Negociaci[oó]n|Presentaci[oó]n)",
    re.IGNORECASE,
)

_METHOD_KEYWORDS = re.compile(
    r"\b(Scrum|Agile|CRISP.?DM|Design\s*Thinking|Lean|Six\s*Sigma|"
    r"PMBOK|PRINCE2|Kanban|DevOps|DataOps|MLOps|"
    r"Regresi[oó]n\s*lineal|Regresi[oó]n\s*log[ií]stica|An[aá]lisis\s*factorial|"
    r"Series\s*de\s*tiempo|An[aá]lisis\s*de\s*componentes|PCA|"
    r"Validaci[oó]n\s*cruzada|Hiperpar[aá]metros|Overfitting|Underfitting)",
    re.IGNORECASE,
)


# ── docx parsing ──────────────────────────────────────────────────────────────

def _cell_text(table, row_idx: int, col_idx: int = 0) -> str:
    try:
        return table.rows[row_idx].cells[col_idx].text.strip()
    except IndexError:
        return ""


def _table_body(table) -> str:
    """Return all non-header cell text joined."""
    parts: list[str] = []
    for row in table.rows[1:]:
        for cell in row.cells:
            txt = cell.text.strip()
            if txt:
                parts.append(txt)
    return "\n".join(parts)


def _find_table_by_header(tables, keyword: str):
    for t in tables:
        header = _cell_text(t, 0).upper()
        if keyword.upper() in header:
            return t
    return None


def _extract_skills_from_text(text: str) -> dict[str, list[str]]:
    """Return categorized skills extracted via regex from raw text."""
    result: dict[str, list[str]] = {
        "tecnologia": [],
        "skill_tecnica": [],
        "herramienta": [],
        "plataforma": [],
        "skill_transversal": [],
        "metodologia": [],
    }
    seen: set[str] = set()

    platforms = {"AWS", "Azure", "GCP", "Google Cloud", "Databricks", "Snowflake",
                 "BigQuery", "Power Apps", "Power Automate", "Power Platform",
                 "MLflow", "Airflow"}
    tools = {"Power BI", "Tableau", "Looker", "QlikSense", "Excel", "KNIME",
             "Docker", "Kubernetes", "Git", "GitHub", "dbt"}
    tech = {"Python", "R", "SQL", "NoSQL", "Spark", "Hadoop", "Kafka",
            "TensorFlow", "PyTorch", "Keras", "Scikit-learn", "MongoDB",
            "PostgreSQL", "MySQL", "Cassandra", "Redis", "Elasticsearch"}

    for m in _TECH_KEYWORDS.finditer(text):
        raw = m.group(0).strip()
        key = raw.lower()
        if key in seen:
            continue
        seen.add(key)
        normalized = _normalize_skill(raw)
        if any(p.lower() in raw.lower() for p in platforms):
            result["plataforma"].append(normalized)
        elif any(t.lower() in raw.lower() for t in tools):
            result["herramienta"].append(normalized)
        elif any(t.lower() in raw.lower() for t in tech):
            result["tecnologia"].append(normalized)
        else:
            result["skill_tecnica"].append(normalized)

    for m in _TRANSVERSAL_KEYWORDS.finditer(text):
        raw = m.group(0).strip()
        key = raw.lower()
        if key in seen:
            continue
        seen.add(key)
        result["skill_transversal"].append(_normalize_skill(raw))

    for m in _METHOD_KEYWORDS.finditer(text):
        raw = m.group(0).strip()
        key = raw.lower()
        if key in seen:
            continue
        seen.add(key)
        result["metodologia"].append(_normalize_skill(raw))

    return {k: v for k, v in result.items() if v}


def _normalize_skill(raw: str) -> str:
    """Title-case and normalize whitespace."""
    return re.sub(r"\s+", " ", raw.strip()).title()


def _normalize_text(text: str) -> str:
    nfkd = unicodedata.normalize("NFKD", text.lower())
    return "".join(c for c in nfkd if not unicodedata.combining(c)).strip()


# ── standard microcurriculum docx (10-table format) ───────────────────────────

def parse_standard_docx(path: Path) -> dict[str, Any] | None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx")
        sys.exit(1)

    doc = Document(str(path))
    tables = doc.tables
    if len(tables) < 6:
        return None

    prog_table = _find_table_by_header(tables, "PROGRAMA ACADÉMICO")
    asig_table = _find_table_by_header(tables, "DENOMINACIÓN")
    desc_table = _find_table_by_header(tables, "DESCRIPCIÓN")
    ra_table = _find_table_by_header(tables, "RESULTADOS DE APRENDIZAJE")
    media_table = _find_table_by_header(tables, "MEDIOS EDUCATIVOS")
    perfil_table = _find_table_by_header(tables, "PERFIL DEL DOCENTE")

    programa = _cell_text(prog_table, 1) if prog_table else ""
    asignatura = _cell_text(asig_table, 1) if asig_table else ""
    descripcion = _table_body(desc_table) if desc_table else ""
    resultados_raw = _table_body(ra_table) if ra_table else ""
    medios_raw = _table_body(media_table) if media_table else ""
    perfil_raw = _table_body(perfil_table) if perfil_table else ""

    full_text = "\n".join([descripcion, resultados_raw, medios_raw, perfil_raw])
    skills_by_type = _extract_skills_from_text(full_text)

    # learning outcomes: lines starting with a verb (Diseñar, Desarrollar, Elaborar…)
    outcomes: list[str] = []
    for line in resultados_raw.splitlines():
        line = line.strip()
        if len(line) > 20 and re.match(r"^[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+", line):
            # Skip section headers
            if not re.match(r"^(Tema|Contenido|Nota|Bibliograf)", line, re.IGNORECASE):
                outcomes.append(line)

    return {
        "programa": programa,
        "asignatura": asignatura,
        "descripcion": descripcion[:1000],
        "resultados_aprendizaje": outcomes[:10],
        "skills_by_type": skills_by_type,
        "source_document": path.name,
    }


# ── criminología docx (paragraph-based format) ────────────────────────────────

def parse_criminologia_docx(path: Path) -> list[dict[str, Any]]:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        sys.exit(1)

    doc = Document(str(path))

    # Extract program name from heading
    programa = "Especialización en Criminología y Victimología"
    for p in doc.paragraphs[:5]:
        if "CRIMINOLOG" in p.text.upper():
            break

    # Extract asignaturas from Table 1 (Resultados de Aprendizaje table)
    ra_table = None
    for t in doc.tables:
        header = _cell_text(t, 0).upper()
        if "RESULTADO" in header and "APRENDIZAJE" in header:
            ra_table = t
            break

    if ra_table is None:
        return []

    # Each pair of rows: RA text + asignatura name (merged cells)
    # Structure: rows alternate between RA statements and subject names
    asignaturas: dict[str, list[str]] = {}
    current_asig = ""
    for row in ra_table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        unique_cells = list(dict.fromkeys(cells))  # deduplicate merged cells
        if not any(unique_cells):
            continue
        # Detect if this row is an asignatura name (short, no period)
        candidate = unique_cells[-1] if unique_cells else ""
        is_asig = (
            len(candidate) > 5
            and not candidate.startswith("RA")
            and not candidate.startswith("•")
            and len(candidate) < 100
        )
        if is_asig and candidate != current_asig:
            current_asig = candidate
            if current_asig not in asignaturas:
                asignaturas[current_asig] = []
        ra_text = unique_cells[0] if unique_cells else ""
        if ra_text.startswith("RA") and current_asig:
            asignaturas[current_asig].append(ra_text)

    # Also extract competencias from paragraphs
    comp_text = ""
    capture = False
    for p in doc.paragraphs:
        if "Competencias" in p.text and "asignatura" in p.text.lower():
            capture = True
        if capture:
            comp_text += p.text + "\n"

    full_text = "\n".join(p.text for p in doc.paragraphs)
    skills_by_type = _extract_skills_from_text(full_text)

    results = []
    for asig, ras in asignaturas.items():
        if not asig:
            continue
        results.append({
            "programa": programa,
            "asignatura": asig,
            "descripcion": "",
            "resultados_aprendizaje": ras[:10],
            "skills_by_type": skills_by_type,
            "source_document": path.name,
        })

    return results


# ── discover and parse all docx ───────────────────────────────────────────────

def discover_docx() -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for docx_path in sorted(STORAGE_DIR.rglob("*.docx")):
        rel = docx_path.relative_to(STORAGE_DIR)
        folder = rel.parts[0] if len(rel.parts) > 1 else ""

        if "Criminolog" in folder:
            sub = parse_criminologia_docx(docx_path)
            records.extend(sub)
        else:
            parsed = parse_standard_docx(docx_path)
            if parsed:
                records.append(parsed)

    return records


# ── DB interaction ─────────────────────────────────────────────────────────────

def _get_db_connection():
    """Return psycopg2 connection using RAILWAY_DATABASE_URL from .env.local or env."""
    env_path = Path(__file__).parent / ".env.local"
    if env_path.exists():
        from dotenv import load_dotenv  # type: ignore
        load_dotenv(env_path, override=True)

    db_url = os.environ.get("RAILWAY_DATABASE_URL") or os.environ.get("DATABASE_URL")
    if not db_url:
        raise RuntimeError(
            "RAILWAY_DATABASE_URL not set. Add it to .env.local:\n"
            "  RAILWAY_DATABASE_URL=postgresql://user:pass@host:port/dbname"
        )

    import psycopg2  # type: ignore
    return psycopg2.connect(db_url)


def fetch_db_microcurriculos() -> list[dict]:
    conn = _get_db_connection()
    try:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT m.id, m.programa, m.asignatura, m.source_document,
                       COUNT(ms.id) AS skill_count
                FROM microcurriculos m
                LEFT JOIN microcurriculo_skills ms ON ms.microcurriculo_id = m.id
                GROUP BY m.id, m.programa, m.asignatura, m.source_document
                ORDER BY m.programa, m.asignatura
            """)
            cols = [d[0] for d in cur.description]
            return [dict(zip(cols, row)) for row in cur.fetchall()]
    finally:
        conn.close()


def fetch_db_especializaciones() -> list[dict]:
    conn = _get_db_connection()
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT id, nombre FROM especializaciones ORDER BY id")
            return [{"id": row[0], "nombre": row[1]} for row in cur.fetchall()]
    finally:
        conn.close()


def _match_especialization_id(programa: str, especializaciones: list[dict]) -> int | None:
    prog_norm = _normalize_text(programa)
    for esp in especializaciones:
        if _normalize_text(esp["nombre"]) in prog_norm or prog_norm in _normalize_text(esp["nombre"]):
            return esp["id"]
    return None


def insert_microcurriculo(
    conn,
    *,
    programa: str,
    asignatura: str,
    descripcion: str,
    resultados: list[str],
    skills_by_type: dict[str, list[str]],
    source_document: str,
    specialization_id: int | None,
) -> int:
    with conn.cursor() as cur:
        cur.execute(
            """
            INSERT INTO microcurriculos
                (programa, asignatura, descripcion, plan_estudios,
                 resultados_aprendizaje, source_document, specialization_id)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            RETURNING id
            """,
            (
                programa,
                asignatura,
                descripcion,
                "",
                "\n".join(resultados),
                source_document,
                specialization_id,
            ),
        )
        micro_id = cur.fetchone()[0]

        skill_rows = []
        for tipo, skills in skills_by_type.items():
            for skill in skills:
                skill_rows.append((micro_id, skill, skill, tipo, None, source_document, None))

        if skill_rows:
            cur.executemany(
                """
                INSERT INTO microcurriculo_skills
                    (microcurriculo_id, skill_original, skill_normalized,
                     tipo_skill, confidence_score, source_document, lineage)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
                """,
                skill_rows,
            )
    return micro_id


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Load microcurriculos from docx to DB")
    mode = parser.add_mutually_exclusive_group()
    mode.add_argument("--dry-run", action="store_true", default=True,
                      help="Preview only, no DB writes (default)")
    mode.add_argument("--execute", action="store_true", default=False,
                      help="Actually insert into production DB")
    parser.add_argument("--verbose", action="store_true")
    args = parser.parse_args()

    execute = args.execute

    print("=" * 60)
    print(f"load_microcurriculos.py  mode={'EXECUTE' if execute else 'DRY-RUN'}")
    print("=" * 60)

    # 1. Parse docx
    print("\n[1/4] Parsing docx files...")
    parsed = discover_docx()
    print(f"  Found {len(parsed)} microcurriculum records from docx")
    for r in parsed:
        total_skills = sum(len(v) for v in r["skills_by_type"].values())
        print(f"  · {r['programa']} / {r['asignatura']} — {total_skills} skills — src: {r['source_document']}")

    # 2. DB state (if possible)
    db_available = False
    db_existing: list[dict] = []
    especializaciones: list[dict] = []
    try:
        print("\n[2/4] Querying DB for existing microcurriculos...")
        db_existing = fetch_db_microcurriculos()
        especializaciones = fetch_db_especializaciones()
        db_available = True
        print(f"  DB has {len(db_existing)} existing microcurriculos")
        print(f"  DB has {len(especializaciones)} especializaciones")
    except Exception as exc:
        print(f"  WARNING: DB not reachable — {exc}")
        print("  Skipping DB comparison. Proceeding with docx-only preview.")

    # 3. Determine what to load
    db_keys: set[str] = set()
    if db_available:
        for row in db_existing:
            db_keys.add(_normalize_text(f"{row['programa']}|{row['asignatura']}"))

    to_insert: list[dict] = []
    already_loaded: list[dict] = []
    for r in parsed:
        key = _normalize_text(f"{r['programa']}|{r['asignatura']}")
        if key in db_keys:
            already_loaded.append(r)
        else:
            to_insert.append(r)

    print(f"\n[3/4] Comparison:")
    print(f"  Already in DB  : {len(already_loaded)}")
    print(f"  To insert      : {len(to_insert)}")

    if args.verbose:
        if already_loaded:
            print("  --- Already loaded ---")
            for r in already_loaded:
                print(f"    ✓ {r['programa']} / {r['asignatura']}")
        if to_insert:
            print("  --- To be inserted ---")
            for r in to_insert:
                print(f"    + {r['programa']} / {r['asignatura']}")

    # 4. Preview / insert
    print(f"\n[4/4] {'INSERTING' if execute else 'PREVIEW (dry-run)'}...")
    for r in to_insert:
        spec_id = _match_especialization_id(r["programa"], especializaciones) if db_available else None
        total_skills = sum(len(v) for v in r["skills_by_type"].values())
        print(f"  {'INSERT' if execute else 'WOULD INSERT'}:")
        print(f"    programa          : {r['programa']}")
        print(f"    asignatura        : {r['asignatura']}")
        print(f"    specialization_id : {spec_id}")
        print(f"    resultados        : {len(r['resultados_aprendizaje'])}")
        print(f"    skills            : {total_skills}")
        if args.verbose:
            for tipo, skills in r["skills_by_type"].items():
                if skills:
                    print(f"      {tipo}: {skills}")
        print()

    if execute and to_insert:
        if not db_available:
            print("ERROR: Cannot execute — DB not reachable.")
            sys.exit(1)
        print("Connecting and inserting...")
        conn = _get_db_connection()
        try:
            conn.autocommit = False
            inserted_ids = []
            for r in to_insert:
                spec_id = _match_especialization_id(r["programa"], especializaciones)
                micro_id = insert_microcurriculo(
                    conn,
                    programa=r["programa"],
                    asignatura=r["asignatura"],
                    descripcion=r["descripcion"],
                    resultados=r["resultados_aprendizaje"],
                    skills_by_type=r["skills_by_type"],
                    source_document=r["source_document"],
                    specialization_id=spec_id,
                )
                inserted_ids.append(micro_id)
                print(f"  ✓ Inserted microcurriculo id={micro_id}: {r['asignatura']}")
            conn.commit()
            print(f"\nDone. Inserted {len(inserted_ids)} microcurriculos. IDs: {inserted_ids}")
        except Exception as exc:
            conn.rollback()
            print(f"ERROR during insert — rolled back: {exc}")
            sys.exit(1)
        finally:
            conn.close()
    elif not to_insert:
        print("  Nothing to insert — all docx already loaded in DB.")
    else:
        print("  Dry-run complete. Run with --execute to perform inserts.")


if __name__ == "__main__":
    main()
