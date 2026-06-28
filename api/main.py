from __future__ import annotations

import asyncio
import logging
import os
import subprocess
import sys
import uuid
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from typing import Any

from fastapi import BackgroundTasks, FastAPI, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse

from api.auth import auth_router
from api.contracts import (
    CareerIntelligenceResponse,
    CriticalProgramPageResponse,
    CurriculumSimulationResponse,
    CurriculumRiskResponse,
    ExecutiveNarrativeResponse,
    ExecutiveObservatoryResponse,
    ForecastSummaryResponse,
    HealthResponse,
    MarketForecastPageResponse,
    ObservatoryStatusResponse,
    PaginatedResponse,
    ProgramIntelligenceItem,
    ProgramIntelligencePageResponse,
    Program,
    ProgramDashboardResponse,
    ProgramSummaryResponse,
    ProgramPageResponse,
    RecommendationV2PageResponse,
    RecommendationExplanationResponse,
    AskObservatoryRequest,
    AskObservatoryResponse,
    SearchResponse,
    UniversityMarketAlignmentResponse,
)
from api.logging import RequestLoggingMiddleware, configure_logging
from api import services


configure_logging(os.getenv("LOG_LEVEL", "INFO"))
logger = logging.getLogger("api.main")


def _fallback_program_response(program_id: int) -> dict[str, Any]:
    program_name = f"Programa {program_id}"
    try:
        program_name = services._safe_program_name(program_id)  # type: ignore[attr-defined]
    except Exception:
        pass
    return {
        "especializacion_id": program_id,
        "nombre_especializacion": program_name,
        "rol": "",
        "total_skills_programa": 0,
        "total_herramientas": 0,
        "total_competencias": 0,
        "total_habilidades_blandas": 0,
        "promedio_match_mercado": 0.0,
        "porcentaje_match": 0.0,
        "max_match_mercado": 0.0,
        "total_empleos_relacionados": 0,
        "skills_cubiertas": 0,
        "skills": [],
        "microcurriculum_context": None,
        "curricular_context_source": "fallback",
        "narrativa_ia": "Análisis curricular pendiente de datos suficientes.",
    }


def _fallback_dashboard_response(program_id: int) -> dict[str, Any]:
    program = _fallback_program_response(program_id)
    return {
        "program_id": program_id,
        "program": program,
        "kpis": {
            "alignment_score": 0.0,
            "missing_critical_skills": 0,
            "high_demand_roles": 0,
            "employability_trend": 0.0,
            "digital_coverage": 0.0,
            "curricular_update_signal": "Baja",
        },
        "status": {
            "curricular_status": "Fallback",
            "curricular_status_detail": "Vista de contingencia mientras se recupera la evidencia.",
            "ai_signal": "Análisis institucional recuperado en modo fallback.",
            "trend_label": "Sin señal suficiente",
        },
        "missing_skills": [],
        "matches": [],
        "recommendations": [],
        "insights": {
            "detected": "No se pudo construir el panel completo; se muestra el programa con evidencia base.",
            "ai_recommends": [],
            "emerging_gap": "Sin brecha prioritaria identificada en el modo fallback",
            "critical_signal": "Fallback institucional",
        },
        "source": "fallback",
    }


def _cors_origins() -> list[str]:
    raw = os.getenv("CORS_ORIGINS") or os.getenv("API_CORS_ORIGINS") or "*"
    if raw.strip() == "*":
        return ["*"]
    return [origin.strip() for origin in raw.split(",") if origin.strip()]


@asynccontextmanager
async def lifespan(_app: FastAPI):
    try:
        snapshot = services.get_health_snapshot()
        logger.info(
            "startup_validation_complete",
            extra={
                "source": "api",
                "database": snapshot.get("database"),
                "status": snapshot.get("status"),
            },
        )
    except Exception:
        logger.exception("startup_validation_failed", extra={"source": "api"})
    yield


class _UnicodeJSONResponse(JSONResponse):
    """Emit Spanish characters as-is instead of \\uXXXX escapes."""
    def render(self, content: Any) -> bytes:
        import json
        return json.dumps(content, ensure_ascii=False, allow_nan=False).encode("utf-8")


app = FastAPI(
    title="AI Labor & Curriculum Observatory API",
    version="1.0.0",
    description="Public API for observatory metrics, recommendations, semantic roles, company intelligence and market forecasts.",
    lifespan=lifespan,
    default_response_class=_UnicodeJSONResponse,
)
app.add_middleware(RequestLoggingMiddleware)
app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins(),
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.include_router(auth_router, prefix="/auth")
app.include_router(auth_router, prefix="/api/auth", include_in_schema=False)


@app.get("/", tags=["system"])
def root() -> dict[str, Any]:
    return {
        "name": "AI Labor & Curriculum Observatory API",
        "status": "ready",
        "docs": "/docs",
        "health": "/health",
        "metrics": "/metrics",
    }


@app.get("/health", response_model=HealthResponse, tags=["system"])
def health() -> dict[str, Any]:
    return services.get_health_snapshot()


@app.get("/api/health", response_model=HealthResponse, tags=["system"], include_in_schema=False)
def api_health() -> dict[str, Any]:
    return services.get_health_snapshot()


@app.get("/readiness", response_model=HealthResponse, tags=["system"])
def readiness() -> dict[str, Any]:
    return services.get_readiness_snapshot()


@app.get("/api/readiness", response_model=HealthResponse, tags=["system"], include_in_schema=False)
def api_readiness() -> dict[str, Any]:
    return services.get_readiness_snapshot()


@app.get("/observatory-status", response_model=ObservatoryStatusResponse, tags=["system"])
def observatory_status() -> dict[str, Any]:
    return services.get_observatory_status()


@app.get("/api/observatory-status", response_model=ObservatoryStatusResponse, tags=["system"], include_in_schema=False)
def api_observatory_status() -> dict[str, Any]:
    return services.get_observatory_status()


@app.get("/liveness", tags=["system"])
def liveness() -> dict[str, Any]:
    return {"status": "alive"}


@app.get("/metrics", response_model=PaginatedResponse, tags=["observatory"])
def metrics(
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
    metric_category: str | None = Query(default=None),
    metric_name: str | None = Query(default=None),
) -> dict[str, Any]:
    return services.list_observatory_metrics(limit=limit, offset=offset, metric_category=metric_category, metric_name=metric_name)


@app.get("/api/programas", response_model=ProgramPageResponse, tags=["programas"])
def programas(
    limit: int = Query(25, ge=1, le=100),
    offset: int = Query(0, ge=0),
) -> dict[str, Any]:
    try:
        return services.list_programas_compatibility(limit=limit, offset=offset)
    except Exception as exc:
        logger.warning("programas_route_fallback: %s", exc, exc_info=True)
        return {"items": [], "count": 0, "total": 0, "limit": limit, "offset": offset}


@app.get("/api/programas/{program_id}", response_model=Program, tags=["programas"])
def programa(program_id: int) -> dict[str, Any]:
    try:
        return services.get_programa_compatibility(program_id)
    except KeyError as exc:
        from fastapi import HTTPException

        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except Exception as exc:
        logger.warning("programa_route_fallback: %s", exc, exc_info=True)
        return _fallback_program_response(program_id)


@app.get("/api/dashboard/programa/{program_id}", response_model=ProgramDashboardResponse, tags=["dashboard"])
def dashboard_programa(program_id: int) -> dict[str, Any]:
    try:
        return services.get_program_dashboard_compatibility(program_id)
    except KeyError as exc:
        from fastapi import HTTPException

        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except Exception as exc:
        logger.warning("dashboard_programa_route_fallback: %s", exc, exc_info=True)
        return _fallback_dashboard_response(program_id)


@app.get("/curriculum-gaps", response_model=PaginatedResponse, tags=["observatory"])
def curriculum_gaps(
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
    specialization: str | None = Query(default=None),
    program_id: int | None = Query(default=None, ge=1),
) -> dict[str, Any]:
    return services.list_curriculum_gaps(limit=limit, offset=offset, specialization=specialization, program_id=program_id)


@app.get("/recommendations", response_model=PaginatedResponse, tags=["observatory"])
def recommendations(
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
    recommendation_type: str | None = Query(default=None),
    target_company: str | None = Query(default=None),
) -> dict[str, Any]:
    return services.list_recommendations(limit=limit, offset=offset, recommendation_type=recommendation_type, target_company=target_company)


@app.get("/recommendations-v2", response_model=RecommendationV2PageResponse, tags=["predictive"])
def recommendations_v2(
    program_id: int | None = Query(default=None),
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
) -> dict[str, Any]:
    return services.list_recommendations_v2(program_id=program_id, limit=limit, offset=offset)


@app.get("/emerging-skills", response_model=PaginatedResponse, tags=["observatory"])
def emerging_skills(
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
) -> dict[str, Any]:
    return services.list_emerging_skills(limit=limit, offset=offset)


@app.get("/semantic-roles", response_model=PaginatedResponse, tags=["observatory"])
def semantic_roles(
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
    role_family: str | None = Query(default=None),
) -> dict[str, Any]:
    return services.list_semantic_roles(limit=limit, offset=offset, role_family=role_family)


@app.get("/company-intelligence", response_model=PaginatedResponse, tags=["observatory"])
def company_intelligence(
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
) -> dict[str, Any]:
    return services.list_company_intelligence(limit=limit, offset=offset)


@app.get("/career-paths", response_model=PaginatedResponse, tags=["observatory"])
def career_paths(
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
) -> dict[str, Any]:
    return services.list_career_paths(limit=limit, offset=offset)


@app.get("/market-forecast", response_model=MarketForecastPageResponse, tags=["observatory"])
def market_forecast(
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
    entity_type: str | None = Query(default=None),
    entity_name: str | None = Query(default=None),
    horizon_months: int | None = Query(default=None, ge=1, le=36),
) -> dict[str, Any]:
    return services.list_market_forecast(limit=limit, offset=offset, entity_type=entity_type, entity_name=entity_name, horizon_months=horizon_months)


@app.get("/programas/{program_id}/curriculum-risk", response_model=CurriculumRiskResponse, tags=["predictive"])
def curriculum_risk(program_id: int) -> dict[str, Any]:
    return services.get_curriculum_risk_index(program_id)


@app.get("/programas/{program_id}/alignment", response_model=UniversityMarketAlignmentResponse, tags=["predictive"])
def curriculum_alignment(program_id: int) -> dict[str, Any]:
    return services.get_university_market_alignment(program_id)


@app.get("/critical-programs", response_model=CriticalProgramPageResponse, tags=["predictive"])
def critical_programs(
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
    horizon_months: int = Query(12, ge=1, le=24),
) -> dict[str, Any]:
    return services.get_critical_programs(limit=limit, offset=offset, horizon_months=horizon_months)


@app.get("/curriculum-simulator", response_model=CurriculumSimulationResponse, tags=["predictive"])
def curriculum_simulator(
    program_id: int = Query(..., ge=1),
    proposed_skills: str | None = Query(default=None),
    horizon_months: int = Query(12, ge=1, le=24),
) -> dict[str, Any]:
    return services.get_curriculum_simulator(program_id, proposed_skills=proposed_skills, horizon_months=horizon_months)


@app.get("/forecast-summary", response_model=ForecastSummaryResponse, tags=["predictive"])
def forecast_summary(limit: int = Query(25, ge=1, le=50)) -> dict[str, Any]:
    return services.get_forecast_summary(limit=limit)


@app.get("/career-intelligence", response_model=CareerIntelligenceResponse, tags=["predictive"])
def career_intelligence(source_role: str | None = Query(default=None), limit: int = Query(12, ge=1, le=25)) -> dict[str, Any]:
    return services.get_career_intelligence(source_role=source_role, limit=limit)


@app.get("/executive-observatory", response_model=ExecutiveObservatoryResponse, tags=["predictive"])
def executive_observatory() -> dict[str, Any]:
    return services.get_executive_observatory()


@app.get("/executive-narrative", response_model=ExecutiveNarrativeResponse, tags=["predictive"])
def executive_narrative(program_id: int | None = Query(default=None, ge=1)) -> dict[str, Any]:
    return services.get_executive_narrative(program_id=program_id)


@app.get("/program-summary/{program_id}", response_model=ProgramSummaryResponse, tags=["predictive"])
def program_summary(program_id: int) -> dict[str, Any]:
    return services.get_program_summary(program_id)


@app.get("/recommendation-explanation/{recommendation_id}", response_model=RecommendationExplanationResponse, tags=["predictive"])
def recommendation_explanation(recommendation_id: int) -> dict[str, Any]:
    return services.get_recommendation_explanation(recommendation_id)


@app.post("/ask-observatory", response_model=AskObservatoryResponse, tags=["predictive"])
def ask_observatory(payload: AskObservatoryRequest) -> dict[str, Any]:
    return services.ask_observatory(
        payload.question,
        program_id=payload.program_id,
        recommendation_id=payload.recommendation_id,
        context=payload.context,
    )


@app.get("/program-intelligence", response_model=ProgramIntelligencePageResponse, tags=["predictive"])
def program_intelligence(
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
) -> dict[str, Any]:
    return services.list_program_intelligence(limit=limit, offset=offset)


@app.get("/program-intelligence/{program_id}", response_model=ProgramIntelligenceItem, tags=["predictive"])
def program_intelligence_detail(program_id: int) -> dict[str, Any]:
    try:
        return services.get_program_intelligence(program_id)
    except KeyError as exc:
        from fastapi import HTTPException

        raise HTTPException(status_code=404, detail=str(exc)) from exc


@app.get("/api/dashboard/summary", tags=["dashboard"])
def dashboard_summary(program_id: int | None = Query(default=None)) -> dict[str, Any]:
    """Return aggregated match data from the latest ml run for the frontend dashboard.
    If program_id is provided, results are filtered to that specialization only.
    """
    try:
        from api.database import fetch_all, fetch_one

        run_row = fetch_one(
            "SELECT id, created_at FROM ml_training_runs "
            "WHERE task_name = 'program_job_match' ORDER BY id DESC LIMIT 1"
        )
        if not run_row:
            return {
                "run_id": None, "fecha": None,
                "programas": [], "top_matches": [],
                "totales": {"matches": 0, "alta": 0, "media": 0, "baja": 0},
            }

        run_id: int = int(run_row["id"])
        fecha: str = run_row["created_at"].strftime("%Y-%m-%d") if run_row.get("created_at") else ""

        # Use literal interpolation for the optional filter to avoid param-count mismatches
        pid_filter = f"AND m.especializacion_id = {int(program_id)}" if program_id else ""
        pid_filter_bare = f"AND especializacion_id = {int(program_id)}" if program_id else ""

        prog_rows = fetch_all(
            f"""
            SELECT
                m.especializacion_id                          AS id,
                COALESCE(e.nombre, m.program_name)            AS nombre,
                COUNT(*)                                      AS matches_total,
                ROUND(AVG(m.score_match)::numeric, 1)        AS score_promedio,
                ROUND(MAX(m.score_match)::numeric, 1)        AS score_maximo,
                COUNT(*) FILTER (WHERE m.relevance_label = 'high')   AS lbl_high,
                COUNT(*) FILTER (WHERE m.relevance_label = 'medium') AS lbl_medium,
                COUNT(*) FILTER (WHERE m.relevance_label = 'low')    AS lbl_low
            FROM ml_program_job_matches m
            LEFT JOIN especializaciones e ON e.id = m.especializacion_id
            WHERE m.run_id = {run_id} {pid_filter}
            GROUP BY m.especializacion_id, e.nombre, m.program_name
            ORDER BY score_maximo DESC
            """,
        )

        programas = [
            {
                "id":             int(r["id"]) if r["id"] else None,
                "nombre":         r["nombre"] or "",
                "matches_total":  int(r["matches_total"]),
                "score_promedio": float(r["score_promedio"] or 0),
                "score_maximo":   float(r["score_maximo"] or 0),
                "labels": {
                    "high":   int(r["lbl_high"]),
                    "medium": int(r["lbl_medium"]),
                    "low":    int(r["lbl_low"]),
                },
            }
            for r in prog_rows
        ]

        top_rows = fetch_all(
            f"""
            SELECT
                COALESCE(e.nombre, m.program_name) AS programa,
                m.job_title                        AS empleo,
                COALESCE(m.company, '')            AS empresa,
                ROUND(m.score_match::numeric, 1)   AS score,
                m.relevance_label                  AS label,
                m.skills_en_comun                  AS skills_en_comun,
                m.skills_faltantes                 AS skills_faltantes
            FROM ml_program_job_matches m
            LEFT JOIN especializaciones e ON e.id = m.especializacion_id
            WHERE m.run_id = {run_id} {pid_filter}
            ORDER BY m.score_match DESC
            LIMIT 30
            """,
        )

        top_matches = [
            {
                "programa":        r["programa"] or "",
                "empleo":          r["empleo"] or "",
                "empresa":         r["empresa"],
                "score":           float(r["score"] or 0),
                "label":           r["label"],
                "skills_en_comun": r["skills_en_comun"] if r["skills_en_comun"] is not None else [],
                "skills_faltantes": r["skills_faltantes"] if r["skills_faltantes"] is not None else [],
            }
            for r in top_rows
        ]

        tot_rows = fetch_all(
            f"SELECT relevance_label, COUNT(*) AS cnt "
            f"FROM ml_program_job_matches "
            f"WHERE run_id = {run_id} {pid_filter_bare} "
            f"GROUP BY relevance_label",
        )
        lbl = {r["relevance_label"]: int(r["cnt"]) for r in tot_rows}

        return {
            "run_id":      run_id,
            "fecha":       fecha,
            "programas":   programas,
            "top_matches": top_matches,
            "totales": {
                "matches": sum(lbl.values()),
                "alta":    lbl.get("high", 0) + lbl.get("high_semantic", 0),
                "media":   lbl.get("medium", 0) + lbl.get("medium_semantic", 0),
                "baja":    lbl.get("low", 0) + lbl.get("low_semantic", 0),
            },
        }
    except Exception as exc:
        logger.error("dashboard_summary error program_id=%s: %s", program_id, exc, exc_info=True)
        return {
            "run_id": None, "fecha": None,
            "programas": [], "top_matches": [],
            "totales": {"matches": 0, "alta": 0, "media": 0, "baja": 0},
            "error": str(exc),
        }


@app.get("/api/programs/related-universities/{program_id}")
def related_universities(program_id: int) -> dict[str, Any]:
    try:
        JOIN_TEMPLATE = """
            LEFT JOIN snies_estadisticas_programa s
              ON s.codigo_snies = m.codigo_snies_programa::INTEGER
             AND s.anio = 2024
        """
        BASE_SELECT = """SELECT m.nombre_ies, m.nombre_programa, m.municipio, m.modalidad,
                                m.nivel_academico, m.creditos, m.duracion, m.periodicidad_admision,
                                COALESCE(s.matriculados, 0) AS matriculados,
                                COALESCE(s.graduados, 0)    AS graduados,
                                COALESCE(s.inscritos, 0)    AS inscritos
                         FROM mineducacion_programas_virtuales m"""
        FILTERS = """AND m.estado_programa ILIKE '%activo%'
                     AND m.modalidad ILIKE '%virtual%'
                     AND m.nivel_academico = 'Posgrado'
                     AND m.nombre_programa ILIKE '%especializ%'
                     ORDER BY COALESCE(s.matriculados, 0) DESC LIMIT 50"""

        PROGRAM_WHERE = {
            94: """WHERE (m.nombre_programa ILIKE '%analytic%'
                      OR m.nombre_programa ILIKE '%datos%'
                      OR m.nombre_programa ILIKE '%big data%'
                      OR m.nombre_programa ILIKE '%inteligencia de negocio%'
                      OR m.nombre_programa ILIKE '%business intelligence%')""",
            92: """WHERE (m.nombre_programa ILIKE '%inteligencia artificial%'
                      OR m.nombre_programa ILIKE '%machine learning%'
                      OR m.nombre_programa ILIKE '%ciencia de datos%'
                      OR m.nombre_programa ILIKE '%data science%')""",
            108: """WHERE (m.nombre_programa ILIKE '%criminolog%'
                       OR m.nombre_programa ILIKE '%forense%'
                       OR m.nombre_programa ILIKE '%criminalistica%'
                       OR m.nombre_programa ILIKE '%seguridad ciudadana%')""",
        }
        where = PROGRAM_WHERE.get(program_id)
        if not where:
            return {"program_id": program_id, "competitors": [], "total": 0}

        sql = f"{BASE_SELECT} {JOIN_TEMPLATE} {where} {FILTERS}"

        from api.database import connection
        with connection() as conn:
            with conn.cursor() as cur:
                cur.execute(sql)
                rows = cur.fetchall()

        competitors = []
        for r in rows:
            competitors.append({
                "nombre_ies":            str(r[0] or ""),
                "nombre_programa":       str(r[1] or ""),
                "ciudad":                str(r[2] or ""),
                "modalidad":             str(r[3] or ""),
                "nivel_academico":       str(r[4] or ""),
                "creditos":              r[5],
                "duracion":              str(r[6] or ""),
                "periodicidad_admision": str(r[7] or ""),
                "matriculados":          int(r[8] or 0),
                "graduados":             int(r[9] or 0),
                "inscritos":             int(r[10] or 0),
            })
        return {"program_id": program_id, "competitors": competitors, "total": len(competitors)}
    except Exception as e:
        logger.error("related_universities error program_id=%s: %s", program_id, e)
        return {"program_id": program_id, "competitors": [], "total": 0, "error": str(e)}


@app.get("/api/dashboard/compare-programs", tags=["dashboard"])
def compare_programs(ids: str = Query(default="94,92,108,20")) -> list[dict]:
    """Compare key metrics across multiple programs for the dashboard Comparativa view."""
    try:
        from api.database import fetch_all

        id_list = [int(x.strip()) for x in ids.split(",") if x.strip().isdigit()]
        if not id_list:
            return []

        rows = fetch_all(
            """
            SELECT
                m.especializacion_id AS id,
                COALESCE(e.nombre, m.program_name) AS nombre,
                ROUND(AVG(m.score_match)::numeric, 1) AS pertinencia,
                COUNT(*) FILTER (
                    WHERE m.skills_en_comun IS NOT NULL
                    AND m.skills_en_comun != '[]'::jsonb
                    AND jsonb_array_length(m.skills_en_comun) > 0
                ) AS empleos_compatibles
            FROM ml_program_job_matches m
            LEFT JOIN especializaciones e ON e.id = m.especializacion_id
            WHERE m.run_id = (
                SELECT MAX(id) FROM ml_training_runs WHERE task_name = 'program_job_match'
            )
              AND m.especializacion_id = ANY(%s)
            GROUP BY m.especializacion_id, e.nombre, m.program_name
            """,
            (id_list,),
        )

        by_id = {int(r["id"]): r for r in rows}
        result = []
        for pid in id_list:
            sa = dashboard_skills_analysis(pid)
            r = by_id.get(pid)
            result.append({
                "id":                  pid,
                "nombre":              (r["nombre"] if r else None) or f"Programa {pid}",
                "pertinencia":         float(r["pertinencia"] or 0) if r else 0.0,
                "cobertura_pct":       sa.get("cobertura_pct", 0.0),
                "empleos_compatibles": int(r["empleos_compatibles"] or 0) if r else 0,
                "brechas_count":       len(sa.get("brechas", [])),
            })

        return sorted(result, key=lambda x: x["pertinencia"], reverse=True)
    except Exception as exc:
        logger.error("compare_programs error: %s", exc, exc_info=True)
        return []


@app.get("/api/dashboard/skills-analysis/{program_id}", tags=["dashboard"])
def dashboard_skills_analysis(program_id: int) -> dict[str, Any]:
    """Bidirectional skills analysis: market demand vs. program curriculum for a given program."""
    import unicodedata as _ud

    def _norm(s: str) -> str:
        """Lowercase + strip accents for cross-set comparison."""
        return "".join(
            c for c in _ud.normalize("NFD", s.lower())
            if _ud.category(c) != "Mn"
        )

    def _match(a: str, b: str) -> bool:
        """Exact or prefix match (min 6 chars) after normalization."""
        na, nb = _norm(a), _norm(b)
        return na == nb or (len(na) >= 6 and (na.startswith(nb) or nb.startswith(na)))

    _EMPTY = {
        "program_id":          program_id,
        "skills_mercado":      [],
        "skills_programa":     [],
        "brechas":             [],
        "fortalezas":          [],
        "exclusivas_programa": [],
        "cobertura_pct":       0.0,
    }
    try:
        from api.database import fetch_all

        # 1. Skills from jobs that actually matched THIS program in the latest run
        market_rows = fetch_all(
            """
            SELECT skill, COUNT(*) AS frecuencia
            FROM (
                SELECT jsonb_array_elements_text(skills_empleo) AS skill
                FROM ml_program_job_matches
                WHERE especializacion_id = %s
                  AND run_id = (
                      SELECT MAX(run_id) FROM ml_program_job_matches
                      WHERE especializacion_id = %s
                  )
                  AND skills_empleo IS NOT NULL
                  AND skills_empleo != '[]'::jsonb
                  AND jsonb_typeof(skills_empleo) = 'array'
                  AND skills_en_comun IS NOT NULL
                  AND skills_en_comun != '[]'::jsonb
                  AND jsonb_array_length(skills_en_comun) > 0
            ) t
            GROUP BY skill
            HAVING COUNT(*) >= 2
            ORDER BY frecuencia DESC
            LIMIT 30
            """,
            (program_id, program_id),
        )
        skills_mercado = [
            {"skill": r["skill"], "frecuencia": int(r["frecuencia"])}
            for r in market_rows
            if r["skill"]
        ]

        # 2. Skills from program curriculum (microcurriculo)
        # Column is skill_normalized (with skill_original as fallback) per load_microcurriculo_excel.py
        prog_rows = fetch_all(
            """
            SELECT COALESCE(ms.skill_normalized, ms.skill_original) AS skill_name,
                   COUNT(*) AS cobertura,
                   ARRAY_REMOVE(ARRAY_AGG(DISTINCT m.asignatura), NULL) AS asignaturas
            FROM microcurriculo_skills ms
            JOIN microcurriculos m ON m.id = ms.microcurriculo_id
            WHERE m.specialization_id = %s
              AND COALESCE(ms.skill_normalized, ms.skill_original) IS NOT NULL
            GROUP BY COALESCE(ms.skill_normalized, ms.skill_original)
            ORDER BY cobertura DESC
            """,
            (program_id,),
        )
        skills_programa = [
            {
                "skill":       r["skill_name"],
                "cobertura":   int(r["cobertura"]),
                "asignaturas": list(r["asignaturas"]) if r["asignaturas"] else [],
            }
            for r in prog_rows
            if r["skill_name"]
        ]

        # 3. Cross analysis — prefix match after normalization
        def _find_prog_match(m_skill: str):
            """Return matching programa skill or None."""
            for p in skills_programa:
                if _match(m_skill, p["skill"]):
                    return p
            return None

        fortalezas = []
        brechas = []
        matched_prog_skills: set[str] = set()

        for m in skills_mercado:
            p = _find_prog_match(m["skill"])
            if p:
                fortalezas.append({
                    "skill": m["skill"],
                    "frecuencia_mercado": m["frecuencia"],
                    "cobertura_programa": p["cobertura"],
                })
                matched_prog_skills.add(p["skill"])
            else:
                brechas.append({"skill": m["skill"], "frecuencia_mercado": m["frecuencia"]})

        exclusivas_programa = [
            {"skill": s["skill"], "cobertura": s["cobertura"]}
            for s in skills_programa
            if s["skill"] not in matched_prog_skills
        ]

        total_mercado = len(skills_mercado)
        cobertura_pct = round(len(fortalezas) / total_mercado * 100, 1) if total_mercado else 0.0

        # Build matriz_completa: union of all skills from both sides
        # Keys: normalized skill name → {demanda_mercado, oferta_programa}
        matriz: dict[str, dict] = {}
        for m in skills_mercado:
            key = _norm(m["skill"])
            matriz[key] = {"skill": m["skill"], "demanda_mercado": m["frecuencia"], "oferta_programa": 0}
        for p in skills_programa:
            key = _norm(p["skill"])
            # Try to find if this prog skill already exists under a market-skill key (prefix match)
            matched_key = next(
                (k for k in matriz if _match(p["skill"], matriz[k]["skill"])), None
            )
            if matched_key:
                matriz[matched_key]["oferta_programa"] = p["cobertura"]
            else:
                matriz[key] = {"skill": p["skill"], "demanda_mercado": 0, "oferta_programa": p["cobertura"]}
        matriz_completa = sorted(matriz.values(), key=lambda x: -x["demanda_mercado"])

        logger.info(
            "skills-analysis/%s: market_skills=%d, prog_skills=%d, fortalezas=%d, brechas=%d, matriz=%d",
            program_id, len(skills_mercado), len(skills_programa), len(fortalezas), len(brechas), len(matriz_completa),
        )

        return {
            "program_id":          program_id,
            "skills_mercado":      skills_mercado,
            "skills_programa":     skills_programa,
            "brechas":             brechas,
            "fortalezas":          fortalezas,
            "exclusivas_programa": exclusivas_programa,
            "cobertura_pct":       cobertura_pct,
            "matriz_completa":     matriz_completa,
        }
    except Exception as exc:
        logger.error("skills_analysis error program_id=%s: %s", program_id, exc, exc_info=True)
        return {**_EMPTY, "error": str(exc)}


@app.get("/semantic-search", response_model=SearchResponse, tags=["search"])
def semantic_search(
    q: str = Query(..., min_length=2, max_length=256),
    entity_type: str = Query(default="job", pattern="^(job|company|skill|role)$"),
    limit: int = Query(10, ge=1, le=25),
) -> dict[str, Any]:
    return services.semantic_search_results(q, entity_type=entity_type, limit=limit)


# ─── Pipeline endpoints ───────────────────────────────────────────────────────

# In-memory job store (Railway restarts clear it — sufficient for UX feedback)
_PIPELINE_JOBS: dict[str, dict[str, Any]] = {}

_SCRIPTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "scripts")
_PYTHON = sys.executable

ACQUISITION_STEPS = [
    ["run_acquisition.py", "--domain", "data_analytics",        "--source", "elempleo", "--limit", "20"],
    ["run_acquisition.py", "--domain", "data_analytics",        "--source", "magneto",  "--limit", "20"],
    ["run_acquisition.py", "--domain", "artificial_intelligence","--source", "elempleo", "--limit", "20"],
    ["run_acquisition.py", "--domain", "criminology",            "--source", "elempleo", "--limit", "20"],
]


def _ts() -> str:
    return datetime.now(timezone.utc).isoformat()


def _run_step(job: dict, label: str, cmd: list[str]) -> bool:
    job["current_step"] = label
    job["log"].append(f"[{_ts()}] START {label}")
    try:
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=600,
            env={**os.environ, "PYTHONUNBUFFERED": "1"},
        )
        out = (result.stdout + result.stderr).strip()
        for line in out.splitlines()[-30:]:     # keep last 30 lines per step
            job["log"].append(line)
        if result.returncode != 0:
            job["log"].append(f"[{_ts()}] ERROR {label} (exit {result.returncode})")
            return False
        job["log"].append(f"[{_ts()}] DONE {label}")
        return True
    except subprocess.TimeoutExpired:
        job["log"].append(f"[{_ts()}] TIMEOUT {label}")
        return False
    except Exception as exc:
        job["log"].append(f"[{_ts()}] EXCEPTION {label}: {exc}")
        return False


def _run_pipeline(job_id: str, steps: list[str], program_id: int | None) -> None:
    job = _PIPELINE_JOBS[job_id]
    job["status"] = "running"
    job["started_at"] = _ts()
    errors: list[str] = []

    try:
        # Step 1 — load microcurrículos from Excel
        if "microcurriculos" in steps:
            ok = _run_step(
                job,
                "microcurriculos",
                [_PYTHON, os.path.join(_SCRIPTS_DIR, "load_microcurriculo_excel.py"), "--execute", "--yes"],
            )
            if not ok:
                errors.append("microcurriculos")

        # Step 2 — acquisition
        if "acquisition" in steps:
            for args in ACQUISITION_STEPS:
                script = args[0]
                label = f"acquisition:{args[2]}:{args[4]}"
                ok = _run_step(
                    job, label,
                    [_PYTHON, os.path.join(_SCRIPTS_DIR, script)] + args[1:],
                )
                if not ok:
                    errors.append(label)

        # Step 3 — matching
        if "matching" in steps:
            match_cmd = [
                _PYTHON,
                os.path.join(_SCRIPTS_DIR, "run_semantic_matching.py"),
                "--persist", "--min-score", "40",
            ]
            if program_id:
                match_cmd += ["--program-id", str(program_id)]
            ok = _run_step(job, "matching", match_cmd)
            if not ok:
                errors.append("matching")

    except Exception as exc:
        job["log"].append(f"[{_ts()}] FATAL: {exc}")
        errors.append("fatal")

    job["status"] = "error" if errors else "done"
    job["finished_at"] = _ts()
    job["errors"] = errors
    job["current_step"] = None


@app.post("/api/pipeline/run", tags=["pipeline"])
def pipeline_run(
    background_tasks: BackgroundTasks,
    body: dict[str, Any] = {},
) -> _UnicodeJSONResponse:
    steps_raw = body.get("steps") or ["microcurriculos"]
    program_id = body.get("program_id") or None
    if program_id is not None:
        try:
            program_id = int(program_id)
        except (TypeError, ValueError):
            program_id = None

    job_id = str(uuid.uuid4())[:8]
    _PIPELINE_JOBS[job_id] = {
        "job_id": job_id,
        "status": "queued",
        "steps": steps_raw,
        "program_id": program_id,
        "created_at": _ts(),
        "started_at": None,
        "finished_at": None,
        "current_step": None,
        "errors": [],
        "log": [],
    }
    background_tasks.add_task(_run_pipeline, job_id, steps_raw, program_id)
    return _UnicodeJSONResponse({
        "job_id": job_id,
        "status": "queued",
        "message": "El matching semántico corre localmente o via GitHub Actions (cada noche). El botón solo recarga microcurrículos desde Excel.",
    })


# ─── Market Context ───────────────────────────────────────────────────────────

_MARKET_STUDIES = [
    {
        "title": "World Economic Forum — Future of Jobs Report 2025",
        "summary": (
            "Cifras clave verificadas: "
            "(1) La IA y la gestión de macrodatos (big data) encabezan la lista de habilidades de más rápido "
            "crecimiento global, seguidas de redes/ciberseguridad y alfabetización tecnológica. "
            "(2) El pensamiento analítico es la habilidad básica más buscada: 7 de cada 10 empresas la "
            "consideran esencial. "
            "(3) Se proyecta una creación neta de 170 millones de empleos para 2030, siendo especialistas "
            "en big data, ingenieros fintech y especialistas en IA/ML los roles de mayor crecimiento "
            "porcentual. "
            "(4) El 39% de las habilidades actuales de los trabajadores se transformará o quedará obsoleta "
            "para 2030. "
            "(5) El 59% de la fuerza laboral mundial necesitará upskilling o reskilling en los próximos años. "
            "(6) El 63% de los empleadores identifica la brecha de habilidades como su principal barrera "
            "de transformación."
        ),
        "url": "https://www.weforum.org/publications/the-future-of-jobs-report-2025",
    },
    {
        "title": "OLE Colombia — Observatorio Laboral para la Educación, Ministerio de Educación Nacional",
        "summary": (
            "Cifras verificadas del sistema oficial de seguimiento a egresados de educación superior en Colombia: "
            "(1) En 2023, los graduados de posgrado registraron una tasa de vinculación laboral (cotizantes) "
            "del 91,3%, frente a 73,4% en pregrado — la formación de posgrado mejora significativamente "
            "la inserción laboral. "
            "(2) La tasa nacional de cotizantes fue de 77,9% en 2022, con tendencia ascendente sostenida "
            "desde 2010."
        ),
        "url": "https://ole.mineducacion.gov.co",
    },
    {
        "title": "Adecco Colombia — Tendencias del Mercado Laboral 2026",
        "summary": (
            "Cifras verificadas sobre el mercado laboral colombiano en 2026: "
            "(1) Las cinco áreas con mayor proyección de empleabilidad son: analítica de datos / business "
            "intelligence / ciencia de datos, ciberseguridad, automatización e integración de IA, ventas "
            "consultivas con enfoque analítico, y sostenibilidad/transición energética. "
            "(2) La selección de talento prioriza evidencia de competencias medibles sobre volumen de "
            "contratación. "
            "(3) El 68% de las empresas en Colombia reporta dificultad para encontrar talento con las "
            "competencias técnicas y digitales requeridas [cifra pendiente de verificación final]."
        ),
        "url": "https://www.adecco.com.co",
    },
]

# In-memory cache: program_id → {analisis, generado_en}
_MARKET_CONTEXT_CACHE: dict[int, dict[str, Any]] = {}


@app.post("/api/dashboard/market-context/{program_id}", tags=["dashboard"])
def market_context(program_id: int, refresh: bool = False) -> dict[str, Any]:
    """Generate AI-personalized market context analysis for a program, cached per program_id.
    Pass ?refresh=true to force regeneration (bypasses in-memory cache).
    """
    import datetime

    # Return cached result if available (unless refresh requested)
    if not refresh and program_id in _MARKET_CONTEXT_CACHE:
        return _MARKET_CONTEXT_CACHE[program_id]

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return {
            "analisis": "Análisis no disponible: OPENAI_API_KEY no configurada.",
            "generado_en": datetime.datetime.utcnow().isoformat(),
            "error": True,
        }

    try:
        # Get program skills data
        sa = dashboard_skills_analysis(program_id)
        brechas    = [b["skill"] for b in sa.get("brechas", [])[:8]]
        fortalezas = [f["skill"] for f in sa.get("fortalezas", [])[:6]]

        # Get program name
        from api.database import fetch_one
        row = fetch_one(
            "SELECT nombre FROM especializaciones WHERE id = %s",
            (program_id,),
        )
        nombre_programa = row["nombre"] if row else f"Programa {program_id}"

        # Get pertinencia score from latest run
        score_row = fetch_one(
            """
            SELECT ROUND(AVG(score_match)::numeric, 1) AS score
            FROM ml_program_job_matches
            WHERE especializacion_id = %s
              AND run_id = (SELECT MAX(id) FROM ml_training_runs WHERE task_name = 'program_job_match')
            """,
            (program_id,),
        )
        score = float(score_row["score"] or 0) if score_row else 0.0

        # Derive domain label from program name for domain-specific framing
        nombre_lower = nombre_programa.lower()
        if any(w in nombre_lower for w in ("dato", "analítica", "analytics", "inteligencia artificial", " ia ", "machine", "big data")):
            dominio = "tecnología de datos e inteligencia artificial"
        elif any(w in nombre_lower for w in ("ciberseguridad", "seguridad inform", "redes", "devops", "sistemas")):
            dominio = "ciberseguridad y tecnología"
        elif any(w in nombre_lower for w in ("crimin", "victimol", "derecho", "juridic", "penal")):
            dominio = "ciencias jurídicas y sociales"
        elif any(w in nombre_lower for w in ("neuropsicol", "psicol", "educaci", "pedagog", "clíni")):
            dominio = "salud mental y educación"
        elif any(w in nombre_lower for w in ("gestion", "gestión", "administr", "negocio", "finanz", "contabl")):
            dominio = "gestión y negocios"
        else:
            dominio = "su área disciplinar"

        # Top skills demanded by market for this program (with frequency)
        skills_mercado_top = sa.get("skills_mercado", [])[:10]
        mercado_str = ", ".join(
            f"{s['skill']} ({s['frecuencia']} vacantes)" for s in skills_mercado_top
        ) if skills_mercado_top else "sin datos disponibles"

        # Brechas = market skills absent from curriculum (most critical)
        brechas_str = ", ".join(brechas[:8]) if brechas else "ninguna identificada"

        # Fortalezas = market skills already covered by curriculum
        fortalezas_str = ", ".join(fortalezas[:5]) if fortalezas else "ninguna identificada"

        studies_text = "\n\n".join(
            f"[{s['title']}]\n{s['summary']}" for s in _MARKET_STUDIES
        )

        prompt = (
            f"Eres un asesor experto en pertinencia curricular para posgrados en {dominio}.\n\n"
            f"DATOS DUROS DE MERCADO (fuentes verificadas — úsalos con cifras exactas, no los parafrasees "
            f"vagamente):\n\n"
            f"{studies_text}\n\n"
            f"DATOS ESPECÍFICOS DEL PROGRAMA '{nombre_programa}':\n"
            f"- Score de pertinencia con el mercado laboral: {score}/100\n"
            f"- Skills más demandadas en vacantes compatibles con este programa: {mercado_str}\n"
            f"- Brechas críticas (habilidades que el mercado pide y el currículo NO cubre): {brechas_str}\n"
            f"- Fortalezas curriculares (habilidades del mercado que el currículo SÍ cubre): {fortalezas_str}\n\n"
            f"INSTRUCCIÓN: Escribe exactamente 3 oraciones en español, dirigidas a un comité curricular "
            f"de {dominio}. Cada oración debe:\n"
            f"1. Conectar UNA brecha específica del programa (de la lista anterior) con UNA cifra concreta "
            f"del WEF, OLE o Adecco — nombra la fuente y la cifra exacta.\n"
            f"2. Ser distinta e imposible de reutilizar para otro programa diferente (el nombre del programa "
            f"y sus skills específicas deben quedar explícitos).\n"
            f"3. Evitar frases vacías como 'es importante', 'resulta crucial', 'en el contexto actual'. "
            f"Ve directo al dato y la implicación curricular concreta.\n"
            f"No añadas introducción, título ni conclusión — solo las 3 oraciones."
        )

        from openai import OpenAI
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=400,
        )
        analisis = response.choices[0].message.content.strip()

        result = {
            "analisis": analisis,
            "generado_en": datetime.datetime.utcnow().isoformat(),
            "error": False,
        }
        _MARKET_CONTEXT_CACHE[program_id] = result
        return result

    except Exception as exc:
        logger.error("market_context error program_id=%s: %s", program_id, exc, exc_info=True)
        return {
            "analisis": f"Error al generar el análisis: {exc}",
            "generado_en": datetime.datetime.utcnow().isoformat(),
            "error": True,
        }


# ─── Curriculum Redesign ──────────────────────────────────────────────────────

# In-memory cache: program_id → redesign result (cleared on new POST)
_REDESIGN_CACHE: dict[int, dict[str, Any]] = {}
# Job tracker for async redesign
_REDESIGN_JOBS: dict[str, dict[str, Any]] = {}


def _compute_tfidf_relevance(text: str, skills: list[str]) -> tuple[float, list[tuple[str, float]]]:
    """
    Return (max_similarity, [(skill, score), ...]) sorted by per-skill relevance to text DESC.

    Uses token overlap against the subject text: each skill token that appears in the text
    contributes to the score. This ensures each subject gets a different ranking because
    the scores depend on the subject's actual vocabulary.
    """
    import unicodedata as _udn
    import math

    def _n(s: str) -> str:
        return "".join(c for c in _udn.normalize("NFD", s.lower()) if _udn.category(c) != "Mn")

    text_norm = _n(text)
    text_tokens = set(text_norm.split())

    ranked: list[tuple[str, float]] = []
    for skill in skills:
        skill_norm = _n(skill)
        skill_tokens = skill_norm.split()
        if not skill_tokens:
            ranked.append((skill, 0.0))
            continue
        # Exact phrase match gets a big bonus
        phrase_bonus = 2.0 if skill_norm in text_norm else 0.0
        # Token overlap: fraction of skill tokens present in text
        hits = sum(1 for t in skill_tokens if t in text_tokens and len(t) >= 4)
        token_score = hits / len(skill_tokens)
        score = phrase_bonus + token_score
        ranked.append((skill, round(score, 4)))

    ranked.sort(key=lambda x: x[1], reverse=True)
    max_score = ranked[0][1] if ranked else 0.0
    return max_score, ranked


def _call_openai_for_ras(asignatura: str, clean_text: str, brechas_relevantes: list[str]) -> dict[str, Any]:
    """Call gpt-4o-mini to propose updated learning outcomes (RAs)."""
    import json as _json

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return {
            "ras_propuestos": [],
            "justificacion": "OPENAI_API_KEY no configurada en el entorno.",
        }

    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key)

        texto_resumido = clean_text[:900] if clean_text else "(sin texto disponible)"
        brechas_str = ", ".join(brechas_relevantes[:8])

        prompt = (
            f"Eres un diseñador curricular universitario colombiano experto en pertinencia académica. "
            f"Aquí están los contenidos/resultados de aprendizaje (RA) actuales "
            f"de la asignatura '{asignatura}':\n\n{texto_resumido}\n\n"
            f"El mercado laboral demanda estas competencias que el programa NO cubre actualmente: "
            f"{brechas_str}\n\n"
            f"INSTRUCCIONES IMPORTANTES:\n"
            f"1. Solo propón RAs si la competencia es académicamente coherente con el campo de '{asignatura}'.\n"
            f"2. Si ninguna de las brechas dadas es coherente con el contenido de la asignatura "
            f"(ej: pedir Excel o habilidades financieras a una asignatura de neuropsicología clínica), "
            f"responde con ras_propuestos: [] y explica en justificacion por qué ninguna aplica.\n"
            f"3. Para las brechas coherentes, propón 1 o 2 RAs NUEVOS o MODIFICADOS en formato académico "
            f"universitario (código RA-XXX, verbo de Bloom, contenido medible y específico).\n\n"
            f"Responde ÚNICAMENTE en JSON con esta estructura exacta:\n"
            f'{{"ras_propuestos": [{{"codigo": "RA-001", "texto": "...", '
            f'"skills_incorporadas": ["..."], "tipo": "nuevo", "ra_original_codigo": null}}], '
            f'"justificacion": "..."}}'
        )

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=0.4,
            max_tokens=800,
        )
        result = _json.loads(response.choices[0].message.content)
        # Ensure expected keys exist
        result.setdefault("ras_propuestos", [])
        result.setdefault("justificacion", "")
        return result
    except Exception as exc:
        logger.error("openai_call_error asignatura=%s: %s", asignatura, exc)
        return {
            "ras_propuestos": [],
            "justificacion": f"Error al llamar a OpenAI: {exc}",
        }


def _run_redesign(job_id: str, program_id: int) -> None:
    """Background task: analyze curriculum vs brechas and call OpenAI for each relevant subject."""
    import unicodedata as _ud

    def _norm(s: str) -> str:
        return "".join(c for c in _ud.normalize("NFD", s.lower()) if _ud.category(c) != "Mn")

    def _match(a: str, b: str) -> bool:
        na, nb = _norm(a), _norm(b)
        return na == nb or (len(na) >= 6 and (na.startswith(nb) or nb.startswith(na)))

    job = _REDESIGN_JOBS[job_id]

    def _update(status: str, **kw: Any) -> None:
        job.update({"status": status, **kw})

    try:
        from api.database import fetch_all

        _update("running", current_step="Cargando microcurrículos")

        # 1. Fetch microcurriculos for this program
        mc_rows = fetch_all(
            """
            SELECT m.id, m.asignatura, m.clean_text,
                   COALESCE(m.asignatura, m.source_document) AS nombre
            FROM microcurriculos m
            WHERE m.specialization_id = %s
              AND m.clean_text IS NOT NULL
              AND length(m.clean_text) > 50
            ORDER BY m.asignatura NULLS LAST
            """,
            (program_id,),
        )

        logger.info("[REDESIGN-DEBUG] program_id=%s mc_rows_count=%d asignaturas=%s",
                    program_id, len(mc_rows),
                    [r.get("nombre") or r.get("asignatura") or str(r.get("id")) for r in mc_rows])

        if not mc_rows:
            _update("error", error="No se encontraron microcurrículos para este programa. Ejecuta primero el pipeline de carga.")
            return

        # 2. Compute brechas using same logic as skills-analysis
        _update("running", current_step="Calculando brechas curriculares")

        market_rows = fetch_all(
            """
            SELECT skill, COUNT(*) AS frecuencia
            FROM (
                SELECT jsonb_array_elements_text(skills_empleo) AS skill
                FROM ml_program_job_matches
                WHERE especializacion_id = %s
                  AND run_id = (SELECT MAX(run_id) FROM ml_program_job_matches WHERE especializacion_id = %s)
                  AND skills_empleo IS NOT NULL AND skills_empleo != '[]'::jsonb
                  AND jsonb_typeof(skills_empleo) = 'array'
            ) t
            GROUP BY skill HAVING COUNT(*) >= 2
            ORDER BY frecuencia DESC LIMIT 30
            """,
            (program_id, program_id),
        )
        skills_mercado = [r["skill"] for r in market_rows if r["skill"]]

        prog_rows = fetch_all(
            """
            SELECT COALESCE(ms.skill_normalized, ms.skill_original) AS skill_name
            FROM microcurriculo_skills ms
            JOIN microcurriculos m ON m.id = ms.microcurriculo_id
            WHERE m.specialization_id = %s
              AND COALESCE(ms.skill_normalized, ms.skill_original) IS NOT NULL
            GROUP BY COALESCE(ms.skill_normalized, ms.skill_original)
            """,
            (program_id,),
        )
        skills_programa = [r["skill_name"] for r in prog_rows if r["skill_name"]]

        # Brechas = market skills not matched by program skills
        brechas: list[str] = []
        for ms in skills_mercado:
            if not any(_match(ms, ps) for ps in skills_programa):
                brechas.append(ms)

        logger.info(
            "[REDESIGN-DEBUG] program_id=%s market_skills=%d prog_skills=%d brechas=%d list=%s",
            program_id, len(skills_mercado), len(skills_programa), len(brechas), brechas,
        )

        if not brechas:
            _update("done", result={
                "program_id": program_id,
                "asignaturas_analizadas": len(mc_rows),
                "advertencia": "No se encontraron brechas curriculares — el programa ya cubre las skills del mercado.",
                "propuestas": [],
            })
            _REDESIGN_CACHE[program_id] = _REDESIGN_JOBS[job_id].get("result", {})
            return

        # FIX 1: Filter generic/noise skills before sending to OpenAI
        # Use token-level matching + accent-stripping so "gestión financiera" matches "financiero"
        import unicodedata as _ud2
        def _norm_noise(s: str) -> str:
            return "".join(c for c in _ud2.normalize("NFD", s.lower()) if _ud2.category(c) != "Mn")

        _RUIDO_EXCLUIDO_TOKENS = {
            'financiero', 'sas', 'excel', 'trabajo en equipo', 'estrategia',
            'gestion', 'comunicacion', 'liderazgo', 'innovacion', 'negociacion',
            'planeacion', 'ventas', 'comercial', 'servicio al cliente', 'contabilidad',
            'presupuesto', 'facturacion', 'logistica', 'power bi', 'tableau',
            'office', 'word', 'powerpoint', 'microsoft', 'erp', 'sap',
        }

        def _is_ruido(skill: str) -> bool:
            normed = _norm_noise(skill)
            # Exact match after normalizing
            if normed in _RUIDO_EXCLUIDO_TOKENS:
                return True
            # Any noise token appears as a word within the skill
            tokens = set(normed.split())
            return bool(tokens & _RUIDO_EXCLUIDO_TOKENS)

        brechas_filtradas = [b for b in brechas if not _is_ruido(b)]
        if not brechas_filtradas:
            brechas_filtradas = brechas  # fallback: use all if filter removes everything

        # 3. Score each asignatura by TF-IDF relevance to brechas
        _update("running", current_step="Calculando relevancia de asignaturas")

        logger.info(
            "[REDESIGN-DEBUG] program_id=%s brechas=%d filtradas=%d brechas_filtradas=%s",
            program_id, len(brechas), len(brechas_filtradas), brechas_filtradas,
        )

        scored: list[dict[str, Any]] = []
        all_scores: list[dict[str, Any]] = []   # kept for debug regardless of threshold
        for row in mc_rows:
            text = (row["clean_text"] or "").strip()
            # FIX 2: get per-skill scores so each subject gets its own ranked brechas
            score, ranked_skills = _compute_tfidf_relevance(text, brechas_filtradas)
            asig_nombre = row["nombre"] or f"Asignatura {row['id']}"
            logger.info(
                "[REDESIGN-DEBUG] asignatura=%r clean_text_len=%d max_score=%.4f top_skills=%s",
                asig_nombre, len(text), score, ranked_skills[:3],
            )
            confianza = "alta" if score > 0.08 else "media" if score > 0.03 else "baja"
            # Top brechas specifically relevant to THIS asignatura
            top_brechas_asig = [skill for skill, _ in ranked_skills[:6]] if ranked_skills else brechas_filtradas[:5]
            all_scores.append({
                "nombre":          asig_nombre[:60],
                "clean_text_len":  len(text),
                "max_score":       round(score, 4),
                "top_brechas":     top_brechas_asig[:4],
            })
            # Always include — no absolute threshold; ranking decides top 5
            scored.append({
                "mc_id":              row["id"],
                "asignatura":         asig_nombre,
                "clean_text":         text,
                "relevancia_score":   round(score, 4),
                "brechas_relevantes": top_brechas_asig,
                "confianza":          confianza,
            })

        scored.sort(key=lambda x: x["relevancia_score"], reverse=True)
        top = scored[:5]

        # Debug block included in every response
        debug_info = {
            "fix_version":        "v4-token-overlap-noise-token-match",
            "brechas_count":      len(brechas),
            "brechas_filtradas_count": len(brechas_filtradas),
            "brechas_sample":     brechas[:10],
            "brechas_filtradas_sample": brechas_filtradas[:10],
            "umbral":             "ninguno (top-5 por ranking de token overlap)",
            "asignaturas_scores": sorted(all_scores, key=lambda x: x["max_score"], reverse=True),
        }

        if not top:
            result = {
                "program_id":             program_id,
                "asignaturas_analizadas": len(mc_rows),
                "advertencia": "Ninguna asignatura superó el umbral de relevancia. Los microcurrículos pueden tener texto insuficiente.",
                "propuestas":             [],
                "debug":                  debug_info,
            }
            _REDESIGN_CACHE[program_id] = result
            _update("done", result=result)
            return

        # 4. Call OpenAI for each top asignatura
        propuestas: list[dict[str, Any]] = []
        for i, item in enumerate(top):
            _update("running", current_step=f"Generando RAs con IA ({i+1}/{len(top)}): {item['asignatura'][:40]}")
            ai_result = _call_openai_for_ras(
                asignatura=item["asignatura"],
                clean_text=item["clean_text"],
                brechas_relevantes=item["brechas_relevantes"],
            )
            # Collect all skills_incorporadas from the proposed RAs
            all_skills: list[str] = []
            for ra in ai_result.get("ras_propuestos", []):
                all_skills.extend(ra.get("skills_incorporadas", []))

            propuestas.append({
                "asignatura":          item["asignatura"],
                "relevancia_score":    item["relevancia_score"],
                "confianza":           item["confianza"],
                "brechas_relevantes":  item["brechas_relevantes"],
                "ras_actuales_texto":  item["clean_text"][:600],
                "ras_propuestos":      ai_result.get("ras_propuestos", []),
                "skills_incorporadas": list(set(all_skills)),
                "justificacion":       ai_result.get("justificacion", ""),
            })

        result = {
            "program_id":             program_id,
            "asignaturas_analizadas": len(mc_rows),
            "brechas_totales":        len(brechas),
            "propuestas":             propuestas,
            "debug":                  debug_info,
        }
        _REDESIGN_CACHE[program_id] = result
        _update("done", result=result)
        logger.info("redesign/%s done: %d proposals", program_id, len(propuestas))

    except Exception as exc:
        logger.error("redesign_error program_id=%s: %s", program_id, exc, exc_info=True)
        _update("error", error=str(exc))


@app.post("/api/curriculum/redesign/{program_id}", tags=["curriculum"])
def curriculum_redesign_start(
    program_id: int,
    background_tasks: BackgroundTasks,
) -> _UnicodeJSONResponse:
    """Start async curriculum redesign analysis. Poll /status/{job_id} for progress."""
    job_id = str(uuid.uuid4())[:8]
    _REDESIGN_JOBS[job_id] = {
        "job_id":       job_id,
        "program_id":   program_id,
        "status":       "queued",
        "current_step": None,
        "error":        None,
        "result":       None,
    }
    background_tasks.add_task(_run_redesign, job_id, program_id)
    return _UnicodeJSONResponse({"job_id": job_id, "status": "queued"})


@app.get("/api/curriculum/redesign/status/{job_id}", tags=["curriculum"])
def curriculum_redesign_status(job_id: str) -> _UnicodeJSONResponse:
    """Poll redesign job status."""
    job = _REDESIGN_JOBS.get(job_id)
    if not job:
        return _UnicodeJSONResponse({"error": "job not found"}, status_code=404)  # type: ignore[name-defined]
    return _UnicodeJSONResponse(job)


@app.get("/api/curriculum/redesign/{program_id}/download", tags=["curriculum"])
def curriculum_redesign_download(program_id: int):
    """Generate and download a .docx report with the redesign proposals."""
    from fastapi.responses import StreamingResponse
    import io

    result = _REDESIGN_CACHE.get(program_id)
    if not result:
        return _UnicodeJSONResponse(
            {"error": "No hay análisis disponible. Ejecuta primero POST /api/curriculum/redesign/{program_id}"},
            status_code=404,  # type: ignore[name-defined]
        )

    propuestas = result.get("propuestas", [])
    if not propuestas:
        return _UnicodeJSONResponse(
            {"error": "El análisis no generó propuestas para este programa.", "advertencia": result.get("advertencia", "")},
            status_code=422,  # type: ignore[name-defined]
        )

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        from api.database import fetch_one
        prog_row = fetch_one("SELECT nombre FROM especializaciones WHERE id = %s", (program_id,))
        prog_nombre = prog_row["nombre"] if prog_row else f"Programa {program_id}"

        title = doc.add_heading(f"Propuesta de Actualización Curricular", level=1)
        title.runs[0].font.color.rgb = RGBColor(0x0D, 0x21, 0x58)

        subtitle = doc.add_paragraph(prog_nombre)
        subtitle.runs[0].bold = True
        subtitle.runs[0].font.size = Pt(13)

        from datetime import date
        doc.add_paragraph(f"Generado: {date.today().isoformat()} · Observatorio UNIR Colombia")
        doc.add_paragraph(
            f"Resumen: {len(propuestas)} asignaturas con propuestas de actualización · "
            f"{result.get('brechas_totales', 0)} brechas identificadas en el mercado."
        )
        doc.add_paragraph("")

        # Per-subject section
        for prop in propuestas:
            doc.add_heading(prop["asignatura"], level=2)
            p = doc.add_paragraph()
            p.add_run("Relevancia con brechas: ").bold = True
            p.add_run(f"{prop['relevancia_score']:.2f} · Brechas relevantes: {', '.join(prop['brechas_relevantes'][:5])}")

            # Comparison table
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "RAs / Contenido Actual"
            hdr[1].text = "RAs Propuestos (IA)"
            for cell in hdr:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x0D, 0x21, 0x58)

            row_cells = table.add_row().cells
            row_cells[0].text = (prop.get("ras_actuales_texto") or "Sin texto disponible")[:500]

            right = row_cells[1]
            for ra in prop.get("ras_propuestos", []):
                p_ra = right.add_paragraph()
                run_code = p_ra.add_run(f"[{ra.get('codigo','RA')}] ")
                run_code.bold = True
                if ra.get("tipo") == "nuevo":
                    run_code.font.color.rgb = RGBColor(0x05, 0x96, 0x69)  # green
                else:
                    run_code.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)  # amber
                p_ra.add_run(ra.get("texto", ""))

            if prop.get("justificacion"):
                jus = doc.add_paragraph()
                jus.add_run("Justificación: ").italic = True
                jus.add_run(prop["justificacion"])

            doc.add_paragraph("")

        # Footer
        doc.add_paragraph(
            "Motor de Pertinencia Académica v2 · UNIR Colombia · Datos: mercado laboral activo Colombia 2024"
        ).runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f"redesign_curricular_programa_{program_id}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as exc:
        logger.error("redesign_download_error program_id=%s: %s", program_id, exc, exc_info=True)
        return _UnicodeJSONResponse({"error": f"Error generando documento: {exc}"}, status_code=500)  # type: ignore[name-defined]


@app.get("/api/pipeline/status/{job_id}", tags=["pipeline"])
def pipeline_status(job_id: str) -> _UnicodeJSONResponse:
    job = _PIPELINE_JOBS.get(job_id)
    if not job:
        return _UnicodeJSONResponse({"error": "job not found"}, status_code=404)  # type: ignore[name-defined]
    # Return last 50 log lines to keep response light
    payload = {**job, "log": job["log"][-50:]}
    return _UnicodeJSONResponse(payload)  # type: ignore[name-defined]
