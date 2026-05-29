from __future__ import annotations

import hashlib
import shutil
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from xml.etree import ElementTree
from pathlib import Path

from microcurriculum_engine.extraction.text_cleaner import normalize_whitespace


STORAGE_DIR = Path("storage/microcurriculos")


@dataclass(frozen=True)
class LoadedDocument:
    source_document: str
    stored_path: str
    filename: str
    extension: str
    content_hash: str
    raw_text: str
    clean_text: str
    extraction_method: str


def content_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def safe_filename(filename: str) -> str:
    base = Path(filename or "microcurriculo.txt").name
    return "".join(ch if ch.isalnum() or ch in ".-_ " else "_" for ch in base).strip() or "microcurriculo.txt"


def store_document_bytes(data: bytes, filename: str, *, storage_dir: Path = STORAGE_DIR) -> Path:
    storage_dir.mkdir(parents=True, exist_ok=True)
    digest = content_hash(data)[:16]
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    target = storage_dir / f"{stamp}_{digest}_{safe_filename(filename)}"
    target.write_bytes(data)
    return target


def copy_document(path: str | Path, *, storage_dir: Path = STORAGE_DIR) -> Path:
    source = Path(path)
    data = source.read_bytes()
    target = store_document_bytes(data, source.name, storage_dir=storage_dir)
    if source.resolve() != target.resolve():
        shutil.copystat(source, target)
    return target


def extract_pdf_text(path: Path) -> tuple[str, str]:
    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        if text.strip():
            return text, "pdfplumber"
    except Exception:
        pass

    try:
        import fitz

        with fitz.open(path) as document:
            text = "\n".join(page.get_text("text") for page in document)
        if text.strip():
            return text, "pymupdf"
    except Exception:
        pass

    return "", "ocr_required"


def extract_docx_text(path: Path) -> tuple[str, str]:
    def from_docx_package() -> str:
        with zipfile.ZipFile(path) as archive:
            xml_names = [
                name
                for name in archive.namelist()
                if name.startswith("word/") and name.endswith(".xml")
                and any(part in name for part in ("document", "header", "footer", "footnotes", "endnotes"))
            ]
            chunks: list[str] = []
            namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for name in xml_names:
                root = ElementTree.fromstring(archive.read(name))
                for node in root.findall(".//w:t", namespace):
                    if node.text:
                        chunks.append(node.text)
                for node in root.findall(".//w:tab", namespace):
                    chunks.append("\t")
                for node in root.findall(".//w:br", namespace):
                    chunks.append("\n")
            return "\n".join(chunk.strip() for chunk in chunks if chunk and chunk.strip())

    try:
        import docx

        document = docx.Document(path)
        chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                chunks.append(" | ".join(cell.text for cell in row.cells if cell.text))
        for section in document.sections:
            chunks.extend(paragraph.text for paragraph in section.header.paragraphs if paragraph.text)
            chunks.extend(paragraph.text for paragraph in section.footer.paragraphs if paragraph.text)
        text = "\n".join(chunk for chunk in chunks if chunk and chunk.strip())
        if text.strip():
            return text, "python-docx"
    except Exception:
        pass

    try:
        text = from_docx_package()
        if text.strip():
            return text, "docx_xml_fallback"
    except Exception:
        pass

    return "", "docx_unavailable"


def extract_text(path: str | Path) -> tuple[str, str]:
    source = Path(path)
    extension = source.suffix.lower()
    if extension == ".pdf":
        return extract_pdf_text(source)
    if extension == ".docx":
        return extract_docx_text(source)
    if extension in {".txt", ".md", ".csv"}:
        return source.read_text(encoding="utf-8", errors="ignore"), "plain_text"
    return source.read_bytes().decode("utf-8", errors="ignore"), "bytes_utf8_fallback"


def load_document(path: str | Path, *, persist_original: bool = True) -> LoadedDocument:
    source = Path(path)
    stored = copy_document(source) if persist_original else source
    raw_text, method = extract_text(stored)
    clean_text = normalize_whitespace(raw_text)
    data = stored.read_bytes()
    return LoadedDocument(
        source_document=str(source),
        stored_path=str(stored),
        filename=source.name,
        extension=source.suffix.lower(),
        content_hash=content_hash(data),
        raw_text=raw_text,
        clean_text=clean_text,
        extraction_method=method,
    )
