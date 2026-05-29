from __future__ import annotations

import re
from pathlib import Path

from docx import Document


OUTPUT_DIR = Path("outputs/rewritten_microcurricula")


def safe_output_name(value: str) -> str:
    replacements = str.maketrans(
        {
            "á": "a",
            "é": "e",
            "í": "i",
            "ó": "o",
            "ú": "u",
            "Á": "A",
            "É": "E",
            "Í": "I",
            "Ó": "O",
            "Ú": "U",
            "ñ": "n",
            "Ñ": "N",
            "¿": "",
            "?": "",
        }
    )
    clean = (value or "microcurriculo").translate(replacements)
    clean = re.sub(r"[^A-Za-z0-9 _.-]+", "_", clean)
    clean = re.sub(r"\s+", " ", clean).strip().replace(" ", "_")
    return clean[:120] or "microcurriculo"


def export_rewritten_docx(
    *,
    document_name: str,
    sections: dict[str, str],
    output_dir: Path = OUTPUT_DIR,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    title = sections.get("DENOMINACIÓN DE LA ASIGNATURA") or Path(document_name).stem
    path = output_dir / f"Microcurriculo_Actualizado_{safe_output_name(title)}.docx"
    doc = Document()
    doc.add_heading("Microcurrículo actualizado", level=1)
    for section, body in sections.items():
        doc.add_heading(section, level=2)
        for paragraph in (body or "Sin información explícita en el documento original.").splitlines():
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    doc.add_paragraph()
    doc.add_paragraph("Nota: créditos, horas, semestre y porcentajes de evaluación se conservan según el documento original.")
    doc.save(path)
    return path


def write_summary_markdown(items: list[dict], output_dir: Path = OUTPUT_DIR) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / "Resumen_Cambios_Visual_Analytics_Big_Data.md"
    lines = [
        "# Resumen De Cambios - Visual Analytics Y Big Data",
        "",
        f"Microcurrículos actualizados: `{len(items)}`",
        "",
    ]
    for item in items:
        lines.extend(
            [
                f"## {item.get('document_name')}",
                "",
                f"- Archivo generado: `{item.get('file_path')}`",
                f"- Cambios trazados: `{len(item.get('change_traceability') or [])}`",
                "",
            ]
        )
    path.write_text("\n".join(lines), encoding="utf-8")
    return path
